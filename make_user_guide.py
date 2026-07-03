"""
Generate: Chassis_Workbench_User_Guide.docx
Covers all 23 tabs of the Motorcycle Chassis Dynamics Workbench (MPAW).
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Styles ────────────────────────────────────────────────────────────────────
styles = doc.styles

def set_style(style_name, font_name='Calibri', size=11, bold=False,
              color=None, space_before=0, space_after=6):
    try:
        st = styles[style_name]
    except KeyError:
        return
    st.font.name = font_name
    st.font.size = Pt(size)
    st.font.bold = bold
    if color:
        st.font.color.rgb = RGBColor(*color)
    st.paragraph_format.space_before = Pt(space_before)
    st.paragraph_format.space_after  = Pt(space_after)

set_style('Heading 1', size=18, bold=True, color=(31,111,235), space_before=12, space_after=6)
set_style('Heading 2', size=14, bold=True, color=(13,110,253), space_before=10, space_after=4)
set_style('Heading 3', size=12, bold=True, color=(52,58,64),   space_before=8,  space_after=3)
set_style('Normal',    size=10.5, space_before=0, space_after=4)

# ── Helpers ───────────────────────────────────────────────────────────────────
def h1(text): return doc.add_heading(text, level=1)
def h2(text): return doc.add_heading(text, level=2)
def h3(text): return doc.add_heading(text, level=3)

def para(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
    p.add_run(text)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run('NOTE: ')
    run.bold = True
    run.font.color.rgb = RGBColor(13, 110, 253)
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor(52, 58, 64)
    r2.font.size = Pt(10)
    return p

def tip(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run('TIP: ')
    run.bold = True
    run.font.color.rgb = RGBColor(25, 135, 84)
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor(52, 58, 64)
    r2.font.size = Pt(10)
    return p

def warn(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run('WARNING: ')
    run.bold = True
    run.font.color.rgb = RGBColor(220, 53, 69)
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor(52, 58, 64)
    r2.font.size = Pt(10)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(60, 60, 60)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D0E4FF')
        cell._tc.get_or_add_tcPr().append(shd)
    for ri, row in enumerate(rows):
        tr = table.rows[ri+1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

def separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_before = Pt(60)
tr = title_para.add_run('Motorcycle Chassis Dynamics Workbench')
tr.bold = True
tr.font.size = Pt(26)
tr.font.color.rgb = RGBColor(31, 111, 235)

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_para.add_run('Complete User Guide — All 23 Tabs')
sr.font.size = Pt(15)
sr.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()
ver_para = doc.add_paragraph()
ver_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
vr = ver_para.add_run('MPAW v1.0  ·  Electron + React + TypeScript\nPhysics: Foale (2002) & Cossalter (2006)')
vr.font.size = Pt(11)
vr.font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════
h1('Table of Contents')
toc_items = [
    ('1',  'Software Overview & Getting Started'),
    ('2',  'Tab 1 — Overview Dashboard'),
    ('3',  'Tab 2 — Geometry'),
    ('4',  'Tab 3 — Mass'),
    ('5',  'Tab 4 — Suspension'),
    ('6',  'Tab 5 — Chain (Drivetrain)'),
    ('7',  'Tab 6 — Ergo (Ergonomics)'),
    ('8',  'Tab 7 — Dynamics'),
    ('9',  'Tab 8 — Graphs'),
    ('10', 'Tab 9 — 3D View'),
    ('11', 'Tab 10 — FEM (Frame Stress)'),
    ('12', 'Tab 11 — Compare'),
    ('13', 'Tab 12 — Simulator (MBD)'),
    ('14', 'Tab 13 — Chassis Sim (Suspension Sweep)'),
    ('15', 'Tab 14 — Anti-Squat'),
    ('16', 'Tab 15 — Chassis Dynamics'),
    ('17', 'Tab 16 — Sweep Compare'),
    ('18', 'Tab 17 — Anti-Dive'),
    ('19', 'Tab 18 — System'),
    ('20', 'Tab 19 — Tire'),
    ('21', 'Tab 20 — Inertia'),
    ('22', 'Tab 21 — Stability'),
    ('23', 'Tab 22 — Fork Compliance'),
    ('24', 'Tab 23 — Aero (Aerodynamics)'),
    ('25', 'Sign Conventions & Coordinate Systems'),
    ('26', 'Quick Reference — Parameter Sensitivity Tables'),
    ('27', 'Export & Data Persistence'),
    ('28', 'Bike Presets (8 Families)'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f'{num}.  ')
    r1.bold = True
    r1.font.size = Pt(10.5)
    r2 = p.add_run(title)
    r2.font.size = Pt(10.5)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — SOFTWARE OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h1('1. Software Overview & Getting Started')

h2('What is MPAW?')
para('The Motorcycle Performance Analysis Workbench (MPAW) is a real-time parametric chassis dynamics tool '
     'built with Electron, React 18, TypeScript, and Zustand state management. It implements validated '
     'physics from Foale "Motorcycle Handling and Chassis Design" and Cossalter "Motorcycle Dynamics" (2006). '
     'Every slider change instantly propagates through 13 physics engines and updates all 23 tabs — no manual '
     'calculation or re-run is needed.')

h2('Primary Use Cases')
for item in [
    'Pre-design geometry studies (trail, rake, anti-squat, swingarm angle)',
    'Suspension tuning analysis (spring rate, sag, natural frequency, damping)',
    'Drivetrain geometry optimisation (sprocket sizing, countershaft position)',
    'Rider ergonomics and centre-of-gravity positioning',
    'Frame structural analysis (FEM: fork, frame tube, swingarm sections)',
    'Aerodynamic assessment and top-speed prediction',
    'Multi-configuration comparison (up to 8 saved configs simultaneously)',
]:
    bullet(item)

h2('How to Start the Application')
para('Development / web mode:')
code_block('cd chassis-workbench && npm run dev   →  open http://localhost:5173')
para('Desktop (Electron):')
code_block('npm run electron:dev    (or launch the built .AppImage / .exe)')

h2('The Interface — Three Areas')
add_table(
    ['Area', 'Description'],
    [
        ['Tab Bar (top)', '23 tabs across the top. Tabs 1–7 show a keyboard shortcut number (1–0). Click any tab to switch. Active tab is highlighted with its accent colour.'],
        ['Left panel', 'Input sliders and controls for the active tab (where applicable). Drag the 5px divider to resize (default 32 % left / 68 % right).'],
        ['Right panel', 'Live 2D engineering drawing (ChassisViz2D) + Results panel — or a full-width analysis dashboard for analysis tabs.'],
    ],
    col_widths=[1.8, 4.5]
)

h2('Family Selector — 8 Bike Presets')
para('The Family dropdown (top-left of the header) loads a complete geometry preset:')
add_table(
    ['Family', 'Rake', 'Trail', 'Key character'],
    [
        ['Sport / Supersport', '24°', '95 mm', '320/220 mm travel, 14/42T sprocket'],
        ['Naked / Roadster',   '25°', '100 mm', '120/130 mm travel, 15/42T'],
        ['Adventure / ADV',    '27°', '120 mm', '200/220 mm travel, 15/42T'],
        ['Cruiser',            '30°', '115 mm', '130/120 mm travel, long swingarm'],
        ['Touring / Luxury',   '30°', '120 mm', '120/110 mm travel, 692 mm swingarm'],
        ['Supermoto',          '25°', '105 mm', '270/180 mm travel, 13/42T, lightweight'],
        ['Enduro / Off-Road',  '26°', '120 mm', '300/280 mm travel, 13/52T, 21" front'],
        ['Scooter / Urban',    '26°', '90 mm',  'CVT mode, 44 mm fork offset'],
    ],
    col_widths=[1.6, 0.7, 0.8, 3.2]
)
note('Selecting a family replaces ALL input parameters. Save your config first (Sweep Compare tab) if needed.')

h2('Header KPI Pills')
para('Four live metrics always visible in the header bar regardless of which tab is open:')
add_table(
    ['Pill', 'Green range', 'What it means'],
    [
        ['Trail (mm)',       '80–120 mm',  'Primary steering feel metric. Below 80 → nervous; above 120 → heavy.'],
        ['Front% (CoG)',     '45–58%',     'Static weight on front tyre as percentage of total.'],
        ['Anti-Squat%',      '80–120%',    'Rear suspension behaviour under acceleration.'],
        ['Safety Factor',    '≥ 3.0',      'Minimum FEM safety factor across all frame elements.'],
    ],
    col_widths=[1.5, 1.3, 3.5]
)

h2('Ground Contact Law (Fundamental Constraint)')
para('Both tyre contact patches always sit at Y = 0 (ground level). The software enforces this automatically: '
     'changing Rear/Front Wheel Diameter instantly updates the corresponding Axle Height. This constraint '
     'cannot be violated and ensures physically correct geometry at all times.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 1 — OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h1('2. Tab 1 — Overview Dashboard')
para('The Overview tab is a read-only dashboard showing a complete health snapshot of the current bike '
     'configuration. No inputs are changed here — go to the relevant tab to adjust parameters.')

h2('Key Performance Indicators (KPI Cards)')
para('Six KPI cards at the top, each showing a live value, a coloured progress bar (green = optimal, '
     'amber = acceptable, red = out-of-range), and a sub-label:')
add_table(
    ['KPI', 'Optimal', 'Sub-info'],
    [
        ['Steering Trail',    '80–120 mm',   'Colour-coded; red below 60 mm or above 150 mm'],
        ['CoG Height',        '—',           'Absolute height from ground + forward position from front axle'],
        ['Weight Split',      '45–58% front','Front / Rear % + axle loads in Newtons'],
        ['Anti-Squat%',       '80–120%',     'Chain contribution % also shown'],
        ['Nat. Freq. Front',  '0.9–1.4 Hz',  'Rear frequency shown in sub-label'],
        ['Min Safety Factor', '≥ 3.0',       'Critical element name shown; "Go to FEM tab" if not solved'],
    ],
    col_widths=[1.8, 1.5, 3.0]
)

h2('Weight Distribution Bar')
para('A horizontal bar split proportionally into Front % (blue) and Rear % (orange), with raw Newton '
     'values for each axle and total mass shown below.')

h2('Bike Profile Radar Chart')
para('Six-axis radar chart scoring the bike 0–100 on:')
add_table(
    ['Axis', 'What it measures', 'Score = 100 when'],
    [
        ['Stability',   'Trail quality',                'Trail 80–120 mm'],
        ['Comfort',     'Front suspension frequency',   'f_n front 0.9–1.4 Hz'],
        ['Traction',    'Anti-Squat quality',           'AS% 65–110%'],
        ['Integrity',   'Frame safety factor',          'Min SF ≥ 5.0'],
        ['Ergonomics',  'Rider knee angle',             'Knee angle 90–130°'],
        ['Balance',     'Weight distribution',          'Front% ≈ 50%'],
    ],
    col_widths=[1.3, 2.5, 2.5]
)

h2('System Health Check')
para('Nine traffic-light indicators (green dot = OK, amber = Check, red = FAIL):')
for item in [
    'Steering Trail — target 80–120 mm',
    'Sag% Front & Rear — target 22–32% of travel',
    'Natural Frequency Front — target 0.9–1.4 Hz',
    'Anti-Squat — target 80–120%',
    'Anti-Dive — higher = less fork dive under braking',
    'Knee Angle — target 90–130°',
    'Hip Angle — target 40–90°',
    'Frame Safety Factor — target ≥ 3.0',
]:
    bullet(item)

h2('Stability Thresholds')
add_table(
    ['Metric', 'Formula', 'Meaning'],
    [
        ['Wheelie Limit (g)', 'g × (WB − X_cg) / Y_cg', 'Acceleration needed to lift front wheel'],
        ['Stoppie Limit (g)', 'g × X_cg / Y_cg',         'Deceleration needed to lift rear wheel'],
        ['Lean Clearance (°)', 'arctan(ground clearance / footpeg offset)', 'Maximum bank angle before footpeg grounds'],
        ['Min Turn Radius (m)', 'WB / sin(lock angle)',   'Tightest U-turn at steering lock'],
        ['Max Grade (°)',      'arctan(μ)',                'Steepest climbable slope at tyre friction μ'],
        ['Load at 0.8 g',     'ΔW = M·0.8g·Y_cg / WB',   'Braking load transfer to front at 0.8 g'],
    ],
    col_widths=[1.8, 2.5, 2.0]
)

h2('Dynamic Load Transfer Row')
para('Computed at the current Dynamics tab scenario values (braking g, accel g, corner speed/radius):')
for item in [
    'Front % Braking — percentage of total weight on front tyre during braking',
    'Front % Accel — percentage of total weight on front tyre during acceleration',
    'Bank Angle — required lean angle for the set corner speed and radius',
    'Lateral Force — centripetal force through CoG at corner',
]:
    bullet(item)

h2('Tire Physics Row')
for item in [
    'Front / Rear Free Radius — unloaded tyre outer radius (mm)',
    'Front / Rear Loaded Radius — under static weight (mm)',
    'Front / Rear Contact Patch Length — critical for traction area estimation (mm)',
    'Natural Frequencies corrected with tyre compliance (series spring model)',
]:
    bullet(item)

h2('Handling Indices Row')
add_table(
    ['Index', 'Formula', 'Interpretation'],
    [
        ['Rear Squat (mm)',      'susp_travel × (1 − AS%/100)',      'How much rear compresses under current accel setting'],
        ['Fork Dive (mm)',       'susp_travel × (1 − AD%/100)',      'How much front compresses under current braking setting'],
        ['Stability Index (SI)', 'trail × WB / 10⁶',                'Higher = more stable straight-line tendency'],
        ['Agility Index (AI)',   'I_yaw / (M × WB²)',               'Lower = more agile, quicker direction change'],
        ['Wobble Sensitivity',   '10⁶ / (trail × WB)',              'Lower = less prone to high-speed wobble'],
        ['Pitch Sensitivity',    'X_cg / WB²  (%/mm)',              'Weight split sensitivity to wheelbase change'],
    ],
    col_widths=[1.8, 2.5, 2.0]
)

h2('Development Progress')
para('Three progress trackers at the bottom of Overview showing completion status of the three development '
     'tracks: MBD Engine (Python), Chassis Sim (Python), and Workbench UI (React/Electron). Each phase shows '
     'a test count badge when tests exist.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 2 — GEOMETRY
# ══════════════════════════════════════════════════════════════════════════════
h1('3. Tab 2 — Geometry')
para('The Geometry tab defines the primary chassis kinematic skeleton. Every other computation derives '
     'from these values. Changes here cascade through all other tabs immediately. The 2D engineering '
     'drawing on the right updates live.')

h2('Section: Steering Geometry')
add_table(
    ['Parameter', 'Range', 'Unit', 'Optimal', 'Effect'],
    [
        ['Head Angle (Rake)', '15–40', '°', '23–28°', 'Steering axis tilt from vertical. More rake = more stable, less agile. Each +1° adds 4–6 mm trail.'],
        ['Fork Offset',       '0–100', 'mm', '25–50 mm', 'Triple-clamp offset from steering axis to axle. Increasing offset REDUCES trail. Primary trail tuning without changing rake.'],
        ['Fork Length',       '400–1000', 'mm', '—', 'Axle to lower triple clamp distance. Sets front axle height and head tube position.'],
        ['Steering Offset',   '−50 to +50', 'mm', '0', 'Fine-tuning offset along the steering axis.'],
    ],
    col_widths=[1.5, 0.8, 0.6, 1.2, 2.2]
)

para('Live result display (blue bar):')
for item in [
    'Trail (mm) — formula: (R_f × cos ε − fork_offset) / sin ε',
    'Mechanical Trail (mm) — trail × cos(rake)',
    'Steering Offset Ground (mm)',
]:
    bullet(item)

h2('Section: Wheel Dimensions')
add_table(
    ['Parameter', 'Range', 'Unit', 'Note'],
    [
        ['Wheelbase', '1200–1800', 'mm', 'Optimal 1380–1480 mm for most bikes. Affects load transfer and stability.'],
        ['Front Wheel Diameter', '500–750', 'mm', 'Automatically sets Front Axle Height = Dia/2 (ground contact law).'],
        ['Rear Wheel Diameter',  '500–750', 'mm', 'Automatically sets Rear Axle Height = Dia/2.'],
        ['Front Axle Height',    '200–500', 'mm', 'Read-only when changed by wheel diameter; can be edited directly.'],
        ['Rear Axle Height',     '200–500', 'mm', 'Same as above. Also affects anti-squat IC position.'],
    ],
    col_widths=[1.8, 1.0, 0.6, 3.0]
)

h2('Section: Swingarm')
add_table(
    ['Parameter', 'Range', 'Unit', 'Optimal', 'Effect'],
    [
        ['Swingarm Length', '350–800', 'mm', '420–580 mm', 'Pivot-to-axle. Longer = better ride quality, more wheelbase, but drops AS% ~3–5% per 40 mm increase.'],
        ['Pivot Height',    '200–500', 'mm', '—', 'Swingarm pivot above ground. Raising it steepens swingarm, raises IC, increases AS%.'],
        ['Pivot X (from front axle)', '600–1200', 'mm', '—', 'Horizontal position of swingarm pivot from front axle.'],
    ],
    col_widths=[1.8, 0.8, 0.6, 1.2, 2.9]
)
para('Live result display: Swingarm Angle (CW+ convention, see Section 25). Typical range +4° to +8°.')

h2('Section: Frame Hardpoints')
add_table(
    ['Parameter', 'Range', 'Unit', 'Note'],
    [
        ['Seat Height',       '600–1000', 'mm', 'Affects rider ergonomics. Typical 780–850 mm road, 900+ ADV.'],
        ['Ground Clearance',  '50–350',   'mm', 'Minimum chassis-to-ground distance. Sets lean limit. Target 120–200 mm road.'],
    ],
    col_widths=[1.8, 1.0, 0.6, 3.0]
)

h2('How Geometry Feeds Other Tabs')
add_table(
    ['Downstream', 'Geometry dependency'],
    [
        ['Trail (all tabs)',    'headAngle, forkOffset, frontWheelDia'],
        ['Anti-Squat%',        'swingarmLength, swingarmPivotHeight, rearAxleHeight, wheelbase'],
        ['Natural Frequency',  'forkLength (via front axle height affecting sprung mass distribution)'],
        ['Anti-Dive%',         'headAngle (rake directly drives AD%)'],
        ['CoG/axle loads',     'wheelbase (denominator of weight split formula)'],
        ['FEM forces',         'geometry drives loads applied to FEM elements'],
    ],
    col_widths=[2.0, 4.3]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 3 — MASS
# ══════════════════════════════════════════════════════════════════════════════
h1('4. Tab 3 — Mass')
para('The Mass tab defines every mass component individually. The physics engine computes the combined '
     'centre-of-gravity (CoG) position and static axle loads from the complete component list. '
     'The 2D diagram shows a coloured dot for each component and a yellow crosshair for the combined CoG.')

h2('Mass Component Fields')
para('Each component row has three fields:')
add_table(
    ['Field', 'Description', 'Convention'],
    [
        ['Mass (kg)', 'Component mass', 'Positive value only'],
        ['X position (mm)', 'Distance from FRONT axle, measured rearward', 'X = 0 at front axle, X = WB at rear axle'],
        ['Y position (mm)', 'Height above ground', 'Y = 0 at ground level'],
    ],
    col_widths=[1.5, 2.8, 2.0]
)

h2('CoG Calculation')
code_block('X_cg = Σ(mᵢ × xᵢ) / Σmᵢ     (mm from front axle, rearward)')
code_block('Y_cg = Σ(mᵢ × yᵢ) / Σmᵢ     (mm above ground)')

h2('Static Axle Loads')
code_block('Front load  = M × g × (WB − X_cg) / WB   (Newtons)')
code_block('Rear load   = M × g × X_cg / WB           (Newtons)')
code_block('Front %     = (WB − X_cg) / WB × 100')

h2('Target Weight Distributions')
add_table(
    ['Bike type', 'Front %', 'Rear %', 'Rationale'],
    [
        ['Sport / Supersport', '52–55%', '45–48%', 'High braking load on front; strong front tyre contact needed'],
        ['Naked / Roadster',   '49–53%', '47–51%', 'Balanced for mixed use'],
        ['Adventure / ADV',    '48–52%', '48–52%', 'Equal split for on/off road versatility'],
        ['Cruiser',            '42–48%', '52–58%', 'Rear-biased; rider seated far back, relaxed geometry'],
    ],
    col_widths=[1.8, 0.9, 0.9, 2.7]
)

h2('Rider Mass Convention')
warn('Rider CoG is at HIP height — approximately seat height + 200 mm above seat surface, NOT at the seat. '
     'All MPAW family presets already use this convention.')

h2('How Mass Affects Other Tabs')
add_table(
    ['Effect', 'Mechanism'],
    [
        ['Anti-Squat % numerator', 'IC height is fixed by geometry; AS% denominator is Y_cg — raising CoG REDUCES AS% numerically'],
        ['Load transfer magnitude', 'ΔW = M × a_g × Y_cg / WB; higher CoG → more weight transfer per g'],
        ['Suspension natural frequency', 'f_n = (1/2π)√(WR/m_sprung); heavier sprung mass lowers frequency'],
        ['FEM element loads', 'Reaction forces on fork and frame scale with total mass and CoG position'],
        ['Inertia moments (I_xx, I_yy, I_zz)', 'Computed from component X/Y positions relative to CoG'],
    ],
    col_widths=[2.2, 4.1]
)

tip('Add the engine, frame, fuel tank, wheels, and rider as separate components for maximum accuracy. '
    'The engine position (typically low and centred) is the most influential single component for CoG height.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 4 — SUSPENSION
# ══════════════════════════════════════════════════════════════════════════════
h1('5. Tab 4 — Suspension')
para('The Suspension tab controls spring rates, damping, travel, and sag for both front fork and rear shock. '
     'It computes wheel rates, natural frequencies, damping ratios, and load transfer at 0.8 g braking.')

h2('Spring Rate vs Wheel Rate')
code_block('Wheel Rate (WR) = Spring Rate × MR²')
code_block('MR = wheel displacement / spring displacement   (≤ 1.0 for linkage bikes)')
para('A progressive linkage (MR < 1) makes the wheel rate always softer than the spring rate. '
     'The wheel rate is what the rider actually feels — it is the effective spring rate at the wheel.')

h2('Front Suspension Parameters')
add_table(
    ['Parameter', 'Range', 'Unit', 'Optimal', 'Note'],
    [
        ['Spring Rate Front', '5–50', 'N/mm', '10–25 N/mm', 'Higher = stiffer fork. Race: 15–25 N/mm; touring: 8–14 N/mm.'],
        ['Damping Coeff Front', '500–8000', 'N·s/m', '—', 'Combined compression + rebound. Compute damping ratio to validate.'],
        ['Fork Travel', '50–350', 'mm', '120–200 mm', 'Total fork stroke available. Sport 120 mm; ADV 200 mm; Enduro 300 mm.'],
        ['Motion Ratio Front', '0.3–1.0', '—', '0.7–0.9', 'Fork is typically direct-acting (MR ≈ 1 for most telescopics).'],
        ['Front Static Sag', '0–100', 'mm', '25–30% travel', 'How much fork compresses under rider weight.'],
    ],
    col_widths=[1.8, 0.8, 0.8, 1.4, 1.6]
)

h2('Rear Suspension Parameters')
add_table(
    ['Parameter', 'Range', 'Unit', 'Optimal', 'Note'],
    [
        ['Spring Rate Rear', '20–150', 'N/mm', '50–100 N/mm', 'Rear spring acts through linkage — wheel rate = k × MR².'],
        ['Damping Coeff Rear', '1000–20000', 'N·s/m', '—', 'Higher = stiffer damping, less compliance.'],
        ['Shock Travel', '50–350', 'mm', '80–150 mm', 'Shock body stroke (not wheel travel). Wheel travel = shock × 1/MR.'],
        ['Motion Ratio Rear', '0.3–1.0', '—', '0.55–0.80', 'Pro-Link / Uni-Trak linkage: typically 0.6–0.75.'],
        ['Rear Static Sag', '0–100', 'mm', '25–30% travel', 'How much shock compresses under rider weight.'],
    ],
    col_widths=[1.8, 0.8, 0.8, 1.4, 1.6]
)

h2('Natural Frequency')
code_block('f_n = (1 / 2π) × √(WR / m_sprung)    [Hz]')
add_table(
    ['Target frequency', 'Front', 'Rear', 'Why front > rear'],
    [
        ['Sport / Supersport', '1.6–2.0 Hz', '1.5–1.9 Hz', 'Prevents pitch resonance coupling (seesaw over bumps)'],
        ['Road / Naked',       '1.3–1.7 Hz', '1.2–1.6 Hz', ''],
        ['Touring',            '1.0–1.4 Hz', '1.0–1.3 Hz', ''],
    ],
    col_widths=[1.8, 1.2, 1.2, 3.0]
)
note('Front natural frequency should always be slightly HIGHER than rear. If front ≈ rear, the bike pitches '
     'rhythmically over undulations (pitch coupling).')

h2('Static Sag')
code_block('Sag% = sag_mm / travel_mm × 100')
para('Target 25–30% of travel for road riding. This places the mid-travel working point equally between '
     'full compression and full extension, giving equal suspension stroke in both directions.')
for item in [
    'Too little sag (< 20%): suspension bottoms out over bumps, harsh ride',
    'Too much sag (> 35%): suspension tops out under acceleration, poor control',
]:
    bullet(item)

h2('Damping Ratio')
code_block('ζ = c / (2 × √(WR × m_sprung))')
add_table(
    ['ζ value', 'Behaviour', 'Typical use'],
    [
        ['< 0.3',   'Underdamped — bouncy, oscillates after bump', 'Compliant off-road'],
        ['0.3–0.5', 'Street/road — controlled but compliant',       'Most road bikes'],
        ['0.5–0.7', 'Sportier — quicker settling',                  'Sport / naked'],
        ['0.7–1.0', 'Overdamped — harsh, reduced grip on rough',    'MotoGP / track'],
    ],
    col_widths=[1.2, 2.5, 2.5]
)

h2('Load Transfer at 0.8 g')
para('Computed and displayed as a summary result: the Newton increase on the front axle at 0.8 g braking '
     '— a standard engineering reference point. Formula: ΔW = M × 0.8g × Y_cg / WB.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 5 — CHAIN
# ══════════════════════════════════════════════════════════════════════════════
h1('6. Tab 5 — Chain (Drivetrain)')
para('The Chain tab defines the drivetrain geometry. This is the most direct input to the anti-squat '
     'Instant Centre (IC) calculation. The 2D diagram shows the countershaft position, both chain runs, '
     'and the Foale IC construction lines in real time.')

h2('Key Parameters')
add_table(
    ['Parameter', 'Range', 'Unit', 'AS% sensitivity', 'Note'],
    [
        ['Front Sprocket (teeth)',   '10–20', 'T', '±3–5% per tooth', 'Countershaft sprocket. Larger radius shifts chain tangent point, moves IC.'],
        ['Rear Sprocket (teeth)',    '30–60', 'T', '±2–4% per tooth (opposite)', 'Larger rear sprocket → steeper chain angle → IC rises.'],
        ['Countershaft X (from front axle)', '500–1100', 'mm', '—', 'Horizontal CS position. Fixed by engine design in production.'],
        ['Countershaft Y (height)', '200–500', 'mm', '+4–7% per 10 mm higher', 'MOST SENSITIVE parameter. CS higher than pivot → steep chain line → IC high.'],
        ['Chain Pitch',             '—', 'mm', '—', '12.7 mm for 530 chain (default). Auto-sets sprocket radii.'],
        ['CVT mode (toggle)',        '—', 'bool', '—', 'For scooters with belt/CVT. Disables chain IC calculation.'],
    ],
    col_widths=[2.0, 0.8, 0.6, 1.4, 1.6]
)

h2('Foale IC Construction — How It Works')
para('The IC (Instant Centre) is the key geometric concept for anti-squat. It is found as the intersection '
     'of two lines:')
for i, item in enumerate([
    'Swingarm extension line — the swingarm axis extended forward past the swingarm pivot',
    'Chain force line — direction of chain tension, extended through the tangent contact point on the DRIVE sprocket rim',
], 1):
    bullet(item, bold_prefix=f'Line {i}: ')

h3('Chain Force Angle (auto-computed)')
code_block('θ_chain = θ_geom + arcsin((r_rear − r_drive) / D)')
para('where D = centre-to-centre distance between sprockets, θ_geom = angle of line joining sprocket centres. '
     'This is the external tangent angle — the direction the chain force actually acts.')

h2('What the 2D Diagram Shows')
add_table(
    ['Element', 'Colour', 'Description'],
    [
        ['Tension run (upper chain)',  'Thick cyan line', 'Physical upper chain path from countershaft to rear sprocket'],
        ['Slack run (lower chain)',    'Thin grey line',  'Lower chain path returning to countershaft'],
        ['Foale force line',           'Dashed orange',   'Chain tension direction extended — this passes through IC'],
        ['Swingarm extension',         'Dashed blue',     'Swingarm axis extended forward and rearward'],
        ['Instant Centre (IC)',        'Yellow dot',      'Intersection of swingarm extension and chain force line'],
    ],
    col_widths=[2.0, 1.5, 2.8]
)

h2('Anti-Squat % from the IC')
code_block('AS% = (IC_height projected to front axle vertical / Y_cg) × 100')
add_table(
    ['AS%', 'Behaviour'],
    [
        ['0%',       'Full squat — rear suspension compresses freely under acceleration'],
        ['50–70%',   'Partial anti-squat — common on naked/ADV bikes'],
        ['80–120%',  'Near-neutral to slight lift — ideal for sport/MotoGP geometry'],
        ['> 120%',   'Jack-up — rear suspension extends under power; spring unloads; traction loss risk'],
    ],
    col_widths=[1.2, 5.1]
)

h2('CVT Mode')
para('Toggle "CVT" on for scooters with belt drive or automatic transmissions. '
     'The chain IC method does not apply (no chain), so the chain force line and IC calculations are suppressed. '
     'Anti-squat percentage is shown as not applicable.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 6 — ERGO
# ══════════════════════════════════════════════════════════════════════════════
h1('7. Tab 6 — Ergo (Ergonomics)')
para('The Ergo tab positions the three rider contact points and two handlebar geometry parameters. '
     'The 2D diagram draws the rider triangle and updates handlebar position live, auto-rotating with '
     'the steering axis when rake angle changes.')

h2('Contact Point Parameters')
add_table(
    ['Contact Point', 'X parameter', 'Y parameter', 'Affects'],
    [
        ['Handlebar grip', 'handlebarX (mm from front axle)', 'handlebarY (mm from ground)', 'Wrist angle, upper body weight, steering torque leverage'],
        ['Seat',           'seatX (mm from front axle)',      'seatY (mm from ground)',      'Rider CoG height and position, hip angle'],
        ['Footpeg',        'footpegX (mm from front axle)',   'footpegY (mm from ground)',   'Knee angle, lower body bracing, lean limit'],
    ],
    col_widths=[1.5, 1.8, 1.8, 2.2]
)

h2('Handlebar Geometry Parameters')
add_table(
    ['Parameter', 'Range', 'Unit', 'Effect'],
    [
        ['Riser Height',       '0–150', 'mm', 'How far the bar rises along the steering axis above the upper triple clamp. Higher riser = more upright posture (ADV / enduro style).'],
        ['Handlebar Reach',    '−200 to +100', 'mm', 'Positive = clip-on (forward, toward front). Negative = pull-back (toward rider, cruiser style).'],
    ],
    col_widths=[1.5, 0.9, 0.6, 3.3]
)
note('The handlebar position is PHYSICALLY DERIVED from the steering geometry. When you change rake angle, '
     'the handlebars automatically rotate with it in the 2D diagram. The Ergo panel displays the result.')

h2('Handlebar Types by Bike Family')
add_table(
    ['Family', 'Type', 'Riser mm', 'Reach mm', 'Visual'],
    [
        ['Sport / Supersport', 'Clip-on', '15', '+40', 'End-on circles + clamp line'],
        ['Naked / Roadster',   'Standard', '55', '−40', 'Single end-on circle'],
        ['ADV',                'Riser',    '90', '−30', 'Wide cross-bar + end circles'],
        ['Cruiser',            'Pull-back', '50', '−130', 'Twin circles + centre bar'],
        ['Touring',            'Pull-back', '45', '−90',  'Twin circles + centre bar'],
        ['Supermoto',          'Wide',     '80', '0',    'Thick bar + end circles'],
        ['Enduro',             'Wide',     '90', '0',    'Thick bar + end circles'],
        ['Scooter',            'Standard', '30', '+50',  'Single end-on circle'],
    ],
    col_widths=[1.7, 1.2, 0.9, 0.9, 1.6]
)

h2('Computed Ergonomic Angles')
add_table(
    ['Angle', 'Target range', 'Effect when out of range'],
    [
        ['Knee angle (°)',  '90–130°', 'Below 90: cramped, fatigue; above 130: poor bracing, heavy steering'],
        ['Hip angle (°)',   '40–90°',  'Below 40: too aggressive; above 90: too upright for control'],
        ['Torso angle (°)', '—',       'Indicates forward lean (clip-on) or upright (cruiser) posture'],
    ],
    col_widths=[1.7, 1.4, 3.2]
)

h2('Why Ergonomics Affects Handling')
for item in [
    'Handlebar height shifts rider CoG — higher bar = more rearward CoG = less front load',
    'Clip-on position (low, forward) shifts CoG forward = more front load + better braking stability',
    'Higher bar gives more leverage to apply steering torque at low speed',
    'Strong forward lean reduces rider exposure to aerodynamic drag',
]:
    bullet(item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 7 — DYNAMICS
# ══════════════════════════════════════════════════════════════════════════════
h1('8. Tab 7 — Dynamics')
para('The Dynamics tab combines three scenario sections (braking, acceleration, cornering) with a full '
     'aerodynamics section. Each section has live load transfer curves that sweep across the full g-range, '
     'not just the single slider point.')

h2('Section: Braking')
add_table(
    ['Parameter / Result', 'Description', 'Target'],
    [
        ['Braking Decel (g)', 'Peak deceleration slider. 1.0 g = hard ABS-level stop.', '0.7–1.0 g'],
        ['Front Load %', 'What fraction of total weight is on the front tyre during braking.', '60–75%'],
        ['ΔW Brake (N)', 'Weight shifted to front = M × a × Y_cg / L.', '—'],
        ['Stoppie g', 'The deceleration at which rear wheel lifts — from Stability engine.', '> 1.0 g ideal'],
    ],
    col_widths=[1.8, 3.0, 1.5]
)

h2('Section: Acceleration')
add_table(
    ['Parameter / Result', 'Description', 'Target'],
    [
        ['Acceleration (g)', 'Peak acceleration slider. 0.3–0.5 g road; > 0.7 g superbike.', '0.2–0.6 g'],
        ['Rear Load %', '100% minus front% under acceleration.', '—'],
        ['ΔW Accel (N)', 'Weight shifted to rear under acceleration.', '—'],
        ['Wheelie g', 'Acceleration at which front wheel lifts (from Stability engine).', '> 0.8 g ideal'],
    ],
    col_widths=[1.8, 3.0, 1.5]
)

h2('Section: Cornering')
add_table(
    ['Parameter / Result', 'Description', 'Target'],
    [
        ['Corner Speed (m/s)', 'Entry or mid-corner speed. Sets centripetal load.', '—'],
        ['Corner Radius (m)', 'Arc radius of the bend. Smaller = sharper corner.', '—'],
        ['Track Width (mm)', 'Lateral tyre contact patch spread.', '50–150 mm'],
        ['Bank Angle (°)', 'Required lean angle. Formula: arctan(V² / (R × g)).', '35–52° optimal'],
        ['Lateral Force (N)', 'Centripetal force through CoG at the set speed/radius.', '—'],
        ['Lean Limit (°)', 'Maximum lean before footpeg contacts ground (from Stability).', '—'],
    ],
    col_widths=[1.8, 3.0, 1.5]
)
para('A bank angle gauge bar visualises the current lean against the 58° physiological limit.')

h2('Load Transfer Curves (live sweep charts)')
para('When data is available these four charts sweep from 0 g to the current slider maximum:')
for item in [
    'Axle Loads vs Decel — front (blue) and rear (red) Newton values vs braking g. Dashed line = current slider.',
    'Anti-Dive % vs Decel — how AD% varies as braking intensity increases. Green band = 20–60% target.',
    'Fork Compression vs Decel — predicted fork travel consumed at each braking level (mm).',
    'Wheelie Margin vs Accel — front axle load percentage as acceleration increases. Drops to zero at wheelie point.',
]:
    bullet(item)

h2('Section: Aerodynamics')
add_table(
    ['Parameter', 'Range', 'Unit', 'Optimal', 'Note'],
    [
        ['Drag Coeff Cx',      '0.15–0.90', '—', '0.30–0.45', 'Faired sport: 0.35–0.40; naked: 0.50–0.65; scooter upright: 0.75+'],
        ['Lift Coeff Cz',      '−0.30 to +0.30', '—', '−0.15 to +0.05', 'Negative = downforce; racing fairing: −0.05 to −0.15'],
        ['Frontal Area',       '0.20–0.80', 'm²', '0.30–0.45', 'Sport: ~0.33 m²; naked: ~0.50 m²; ADV with luggage: ~0.65 m²'],
        ['Pressure Centre X',  '400–1200', 'mm', '—', 'Aerodynamic centre from front axle. Forward of CoG → nose-down aero moment'],
        ['Reference Speed',    '50–350', 'km/h', '—', 'Speed at which single-point KPI output is computed'],
        ['Engine Power',       '20–300', 'kW', '—', 'Used for top-speed prediction only'],
        ['Drivetrain η',       '0.75–0.98', '—', '0.85–0.93', 'Chain/belt efficiency. Chain loss ≈ 10–12%'],
    ],
    col_widths=[1.7, 1.2, 0.6, 1.2, 1.7]
)
para('Aerodynamic results output:')
for item in [
    'Drag force at reference speed (N)',
    'Lift / downforce at reference speed (N)',
    'Drag power consumption at reference speed (kW)',
    'Aero pitch moment (N·m) — front load change due to aero',
    'Top speed prediction (km/h) using: V_max = ∛(2Pη / ρCxA)',
    'Drag & Lift Force vs Speed chart (sweep to max speed)',
    'Drag Power vs Speed chart',
]:
    bullet(item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 8 — GRAPHS
# ══════════════════════════════════════════════════════════════════════════════
h1('9. Tab 8 — Graphs')
para('The Graphs tab is a visual dashboard of all computed results across every category. It provides '
     'a quick overall health-check of the current configuration without needing to navigate between tabs. '
     'No inputs can be changed here.')

h2('Radar Chart')
para('Six-axis radar (spider) chart. Each axis is scored 0–100:')
for item in [
    'Trail — quality of steering geometry (centred on 100 mm)',
    'CoG Balance — weight distribution quality',
    'Anti-Squat Index — drivetrain geometry quality',
    'Suspension Frequency — ride comfort quality',
    'Frame Safety Factor — structural margin quality',
    'Aerodynamic Efficiency — drag coefficient quality',
]:
    bullet(item)
para('A well-balanced all-round bike will show a roughly equal hexagon. Outliers flag problem areas immediately.')

h2('Bar Charts')
para('Category-by-category bar charts comparing current values against target ranges:')
for item in [
    'Green bar = within optimal target range',
    'Amber bar = acceptable but outside ideal',
    'Red bar = outside acceptable range — action recommended',
]:
    bullet(item)

h2('Workflow')
for i, item in enumerate([
    'Set up baseline configuration in Geometry, Mass, and Suspension tabs',
    'Switch to Graphs to get the overall picture',
    'Identify the weakest axes on the radar',
    'Return to the relevant tab, adjust, and come back to confirm improvement',
], 1):
    bullet(item, bold_prefix=f'Step {i}: ')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 9 — 3D
# ══════════════════════════════════════════════════════════════════════════════
h1('10. Tab 9 — 3D View')
para('The 3D tab renders a parametric Three.js 3D model of the chassis that updates live with all '
     'geometry slider changes. It provides an intuitive visual check of the proportions and geometry '
     'before any physical prototype is built.')

h2('Controls')
for item in [
    'Left-click + drag — orbit (rotate view)',
    'Right-click + drag — pan',
    'Scroll wheel — zoom in / out',
    'The model updates automatically on any geometry or dimension change',
]:
    bullet(item)

h2('What is Rendered')
for item in [
    'Front fork tubes with the correct rake angle and fork offset',
    'Head tube cylinder in the correct position',
    'Swingarm at the correct length and angle from pivot to rear axle',
    'Both wheel discs at correct wheel diameters and axle heights',
    'Frame spine connecting head tube to swingarm pivot',
    'Rider silhouette at the ergo contact-point positions',
]:
    bullet(item)

tip('Use the 3D tab to sanity-check proportions after major geometry changes — especially after altering '
    'wheelbase, swingarm length, or wheel diameters. The 2D diagram is more precise; 3D gives intuition.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 10 — FEM
# ══════════════════════════════════════════════════════════════════════════════
h1('11. Tab 10 — FEM (Frame Structural Analysis)')
para('The FEM tab applies a simplified finite-element model to three key structural members: the front fork, '
     'the frame main tube, and the swingarm. It computes bending stress, combined stress, and safety factors '
     'for each element under the current load scenario. The right panel renders a colour-coded 3D stress visualization.')

h2('Input: Section Geometry')
add_table(
    ['Parameter', 'Range', 'Unit', 'Note'],
    [
        ['Fork OD',           '25–60',  'mm', 'Outer diameter of fork tube. Sport: 41–50 mm; MX: 48–50 mm'],
        ['Fork Wall',         '1–8',    'mm', 'Wall thickness of fork tube'],
        ['Frame OD',          '20–50',  'mm', 'Main frame tube outer diameter'],
        ['Frame Wall',        '1–6',    'mm', 'Main frame tube wall thickness'],
        ['Swingarm OD',       '25–60',  'mm', 'Swingarm tube outer diameter'],
        ['Swingarm Wall',     '1–6',    'mm', 'Swingarm tube wall thickness'],
    ],
    col_widths=[1.7, 0.8, 0.6, 3.2]
)

h2('Input: Material Selection')
add_table(
    ['Material', 'Yield Strength', 'Elastic Modulus', 'Typical use'],
    [
        ['Steel (STEEL)',         '250 MPa', '200 GPa', 'Standard chromoly frames, swingarms'],
        ['Aluminium (ALUMINUM)',  '270 MPa', '70 GPa',  'Sport/race frames, fork tubes, swingarms'],
        ['CFRP',                  '600 MPa', '150 GPa', 'Race monocoques, high-end subframes'],
        ['Titanium (TITANIUM)',   '900 MPa', '110 GPa', 'Premium frames, lightweight performance'],
    ],
    col_widths=[1.7, 1.3, 1.4, 2.0]
)

h2('Results')
add_table(
    ['Result', 'Description', 'Target'],
    [
        ['Max Displacement (mm)', 'Maximum node displacement under the current load case', '< 1.0 mm for chassis-critical members'],
        ['Critical Element', 'The element with the lowest safety factor', 'Should not be the fork or main frame'],
        ['σ (MPa) per element', 'Combined bending + axial stress', '< yield / safety_factor'],
        ['Safety Factor (SF)', 'Yield strength / combined stress. Red if SF < 1.5', '≥ 3.0 ideal; ≥ 1.5 minimum'],
    ],
    col_widths=[1.8, 2.8, 1.7]
)

warn('A safety factor below 1.5 is flagged with a ⚠ symbol. Below 1.0 means the element has exceeded '
     'yield strength under the computed loads — NEVER build to this specification.')

h2('FEM 3D Visualization (right panel)')
para('The right pane (FEMScene) renders a Three.js model of the three structural elements, colour-coded '
     'by stress level: green (safe) → yellow (marginal) → red (high stress / danger). Hover over any '
     'element to see its exact stress value and safety factor.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 11 — COMPARE
# ══════════════════════════════════════════════════════════════════════════════
h1('12. Tab 11 — Compare')
para('The Compare tab allows side-by-side comparison of two complete bike configurations: '
     'Bike A (the current store configuration) and Bike B (any preset family or custom bike). '
     'Results are shown as a radar chart, a bar chart, and a numerical results table.')

h2('Selecting Bike B')
para('Use the Bike B dropdown to select from:')
for item in [
    'Any of the 8 preset families (Sport, Naked, ADV, Cruiser, Touring, Supermoto, Enduro, Scooter)',
    'Any custom bike you have saved to the custom bike library',
]:
    bullet(item)

h2('Radar Comparison Chart')
para('Both bikes plotted on the same 6-axis radar. Bike A shown in blue, Bike B in orange. '
     'The overlap area shows where the configurations are similar; gaps show where they diverge.')
para('Axes scored 0–100: Stability (trail), Comfort (suspension freq), Traction (AS%), '
     'Balance (weight split), Ergonomics (knee angle), Integrity (FEM safety factor).')

h2('Bar Chart Comparison')
para('Side-by-side horizontal bars for each metric: Trail, Front%, AS%, Natural Frequency Front, '
     'Knee Angle, Safety Factor. Green target band shown on each bar.')

h2('Numerical Results Table')
add_table(
    ['Metric', 'Bike A', 'Bike B', 'Δ (B − A)'],
    [
        ['Trail (mm)', '—', '—', 'Highlighted green if improvement'],
        ['Rake (°)', '—', '—', ''],
        ['Front % (static)', '—', '—', ''],
        ['Anti-Squat%', '—', '—', ''],
        ['Nat Freq Front (Hz)', '—', '—', ''],
        ['Nat Freq Rear (Hz)', '—', '—', ''],
        ['Knee Angle (°)', '—', '—', ''],
        ['Min Safety Factor', '—', '—', ''],
    ],
    col_widths=[2.0, 1.5, 1.5, 1.3]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 12 — SIMULATOR (MBD)
# ══════════════════════════════════════════════════════════════════════════════
h1('13. Tab 12 — Simulator (MBD)')
para('The Simulator tab is a multi-body dynamics validation dashboard. It does NOT run a live simulation '
     'of the current bike — instead it displays pre-validated MBD engine test results and charts, '
     'demonstrating the correctness of the underlying Python MBD physics engine.')

h2('What is Shown')
para('Three validation charts from the Python MBD engine (Phases 1–3, 32/32 tests passing):')
add_table(
    ['Chart', 'Description'],
    [
        ['Energy Drift — RK4 vs Gen-α', 'Comparison of energy conservation between the classic 4th-order Runge-Kutta integrator and the Generalized-α (Gen-α) DAE solver. Gen-α shows superior stability with < 0.005% energy drift vs RK4 ≈ 0.02%.'],
        ['Constraint Violation (2D vs 3D)', 'Baumgarte stabilization keeps constraint violations below 2 mm for 2D pendulum and < 2.1 mm for 3D with quaternion coupling.'],
        ['Gyroscope Precession', 'Simulated precession rate vs theoretical Ωp = τ/(I₃·ω_spin). Simulation error < 5% across 3 seconds.'],
    ],
    col_widths=[2.2, 4.1]
)

h2('MBD Engine Status')
add_table(
    ['Phase', 'Status', 'Tests'],
    [
        ['Phase 1 · 2D Rigid Body + RK4',    'DONE', '14/14'],
        ['Phase 2 · 3D Joints + Quaternions', 'DONE', '23/23'],
        ['Phase 3 · Gen-α DAE Solver',        'DONE', '32/32'],
        ['Phase 4 · Contact + Collision',      'Next', '—'],
        ['Phase 5 · Flexible Bodies (FEM)',    'Pending', '—'],
    ],
    col_widths=[2.5, 1.2, 1.0]
)
note('This tab will evolve to become a live multi-body motorcycle simulation when Phase 4+ are complete. '
     'Currently it serves as a physics validation reference.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 13 — CHASSIS SIM
# ══════════════════════════════════════════════════════════════════════════════
h1('14. Tab 13 — Chassis Sim (Suspension Sweep)')
para('The Chassis Sim tab sweeps the rear suspension through its full travel range (0 mm to max shock travel) '
     'and plots how key metrics change at every point. This reveals the behaviour of the suspension system '
     'across its entire operating range, not just at the static sag point.')

h2('Left Panel: Shock Mount Geometry Inputs')
add_table(
    ['Parameter', 'Description'],
    [
        ['Linkage Type', 'Select Direct Monoshock or 4-Bar (Pro-Link / Uni-Trak) linkage. 4-Bar uses four additional dimensions to define the leverage curve.'],
        ['Shock Lower Mount — swingarm offset X/Y', 'Position of shock lower mount relative to swingarm pivot. Determines leverage at each travel point.'],
        ['Shock Upper Mount — frame offset X/Y', 'Position of shock upper mount on the frame. Fixed in normal operation.'],
        ['4-Bar link dimensions (if 4-Bar selected)', 'Rocker arm geometry that creates the progressive motion ratio curve.'],
    ],
    col_widths=[2.2, 4.1]
)

h2('Static Position Summary Cards')
add_table(
    ['Card', 'Target', 'Meaning'],
    [
        ['Motion Ratio (MR)', '0.55–0.80', 'Wheel displacement / shock displacement at static sag point'],
        ['Wheel Rate (WR)',   '15–80 N/mm', 'Effective spring rate at the wheel at sag point'],
        ['Anti-Squat%',      '60–110%',    'AS% at static sag position'],
    ],
    col_widths=[1.8, 1.3, 3.2]
)

h2('Suspension Sweep Charts')
add_table(
    ['Chart', 'X axis', 'Y axis', 'What to look for'],
    [
        ['Motion Ratio vs Travel',    'Travel (mm)', 'MR', 'Should increase (become more progressive) toward full compression'],
        ['Wheel Rate vs Travel',      'Travel (mm)', 'WR (N/mm)', 'Progressive WR curve prevents bottoming; flat = linear; rising = progressive'],
        ['Anti-Squat% vs Travel',     'Travel (mm)', 'AS%', 'How AS% changes as rear compresses — important for cornering exit'],
        ['Trail vs Travel',           'Travel (mm)', 'Trail (mm)', 'Trail changes slightly as fork compresses — affects steering feel in corners'],
        ['Swingarm Angle vs Travel',  'Travel (mm)', 'SA angle (°)', 'How swingarm angle changes across travel range'],
    ],
    col_widths=[2.0, 1.2, 1.3, 1.8]
)

h2('Dynamics Sweep Charts')
para('Below the suspension sweep, a set of dynamics charts sweep across braking/acceleration levels:')
for item in [
    'Weight Transfer vs Decel — total N transferred to front as braking increases',
    'Axle Loads vs Decel — front and rear axle loads in N across 0–max braking',
    'Anti-Dive% vs Decel — AD% variation with braking intensity',
    'Fork Compression vs Decel — predicted fork travel consumed vs braking g',
    'Rear Extension vs Decel — rear suspension travel under braking (negative = lift)',
    'Weight Transfer vs Accel — total N transferred to rear as acceleration increases',
    'Wheelie Margin vs Accel — front axle load % as acceleration increases',
]:
    bullet(item)

h2('Export')
para('CSV download buttons export all sweep data for import into Excel or Python for further analysis.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 14 — ANTI-SQUAT
# ══════════════════════════════════════════════════════════════════════════════
h1('15. Tab 14 — Anti-Squat')
para('The Anti-Squat tab is the primary chassis tuning tool for acceleration behaviour. It is a self-contained '
     'interactive playground: changes here do NOT modify the main store — they are local overrides for '
     'instant sensitivity analysis. The baseline always reflects the current store configuration.')

h2('How to Use')
for i, item in enumerate([
    'Select a parameter from the left dropdown (9 available)',
    'Use the delta slider to shift that parameter from its baseline value',
    'Watch the IC diagram, AS%, and Cossalter R update in real time',
    'The sweep chart shows the full parameter range — your delta is marked with a dashed line',
    'Reset to baseline with the Reset button',
], 1):
    bullet(item, bold_prefix=f'Step {i}: ')

h2('The Nine Parameters')
add_table(
    ['#', 'Parameter', 'AS% sensitivity', 'Physical meaning'],
    [
        ['1', 'Rear Axle Height',    '+3–5% per 10 mm higher',    'Raising axle (larger wheel or ride height) raises IC, increases AS%'],
        ['2', 'CSP Height',          '+4–7% per 10 mm higher',    'MOST SENSITIVE. Countershaft position Y above ground'],
        ['3', 'CSP Position X',      'Moderate',                  'CS forward/rearward — shifts IC horizontally'],
        ['4', 'Swingarm Angle (Δ°)', '+4–6% per 1° CW steeper',  'Positive delta = CCW (raises rear) = reduces CW slope = lowers AS%'],
        ['5', 'Swingarm Length',     '−3–5% per 40 mm longer',    'Longer swingarm = shallower chain angle, more wheelbase, lower AS%'],
        ['6', 'Wheelbase',           '−2–4% per 100 mm longer',   'Longer WB = shallower load transfer line slope'],
        ['7', 'CoG Height',          '−2–4% per 50 mm higher',    'Higher CoG raises Y_cg denominator, reducing AS% numerically'],
        ['8', 'Front Sprocket (teeth)', '±3–5% per tooth',        'Larger radius shifts chain tangent — IC moves'],
        ['9', 'Rear Sprocket (teeth)',  '∓2–4% per tooth',        'Opposite effect to front sprocket'],
    ],
    col_widths=[0.3, 1.7, 1.5, 2.8]
)

h2('IC Geometry Diagram')
para('The diagram shows (front wheel on LEFT, rear wheel on RIGHT):')
add_table(
    ['Element', 'Colour', 'Description'],
    [
        ['Swingarm',             'Blue solid',    'Actual swingarm from pivot to rear axle'],
        ['Swingarm extension',   'Blue dashed',   'Extended swingarm axis line forward and rearward'],
        ['Chain force line',     'Orange dashed', 'Direction of chain tension extended past both sprockets'],
        ['Instant Centre (IC)',  'Yellow dot',    'Intersection — the key point for AS% computation'],
        ['Load transfer line',   'Purple',        'From rear contact patch through CoG — defines τ and σ angles'],
        ['σ/τ labels',          'White text',    'Angles at rear contact patch used in Cossalter R formula'],
    ],
    col_widths=[1.8, 1.5, 3.0]
)

h2('Cossalter Squat Ratio R')
code_block('R = tan(τ) / tan(σ)')
para('Where σ = angle of load transfer line, τ = angle of swingarm extension line (both from rear contact patch). '
     'When R = 1.0: 100% AS, neutral suspension. R > 1: jack-up. R < 1: squat.')

h2('Cause-Effect Chain Display')
para('When a parameter is selected, the panel shows a step-by-step cause-effect chain. Example for CSP Height:')
code_block('INCREASE CSP height (+10 mm)\n→ Chain line steepens\n  → Chain force line rotates upward\n    → IC height increases ~10–15 mm\n      → AS% increases +4–7%\n        → Above 100%: rear jack-up under power\n          → Spring unloads → reduced tyre compliance')

h2('Sweep Chart')
para('A sweep chart plots AS% (Y-axis) across the full range of the selected parameter (X-axis). '
     'A green band marks the 80–120% target. Your current delta position is shown as a dashed vertical line. '
     'This immediately shows how much margin exists in either direction.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 15 — CHASSIS DYNAMICS
# ══════════════════════════════════════════════════════════════════════════════
h1('16. Tab 15 — Chassis Dynamics')
para('The Chassis Dynamics tab implements Cossalter (2014) §1.1–1.5 chassis dynamics formulas. '
     'It computes frequency-domain handling modes (capsize, weave, wobble), fork equivalent stiffness, '
     '2-DOF stability modes, and plots sweep charts vs speed and steering angle.')

h2('Local Panel Inputs')
add_table(
    ['Input', 'Range', 'Unit', 'Effect'],
    [
        ['Steering Angle δ',  '1–45', '°', 'Handlebar input angle for steady-state steering analysis'],
        ['Speed V',          '10–250', 'km/h', 'Reference speed for single-point results'],
        ['Corner Radius',    '5–300', 'm', 'Set independently from Dynamics tab'],
    ],
    col_widths=[1.8, 0.8, 0.6, 3.1]
)

h2('Left Panel: Computed Results')
add_table(
    ['Result group', 'Metrics shown'],
    [
        ['Fork Equivalent Stiffness', 'k_fork_eq (N/mm) — combined lateral stiffness at current speed and steer angle'],
        ['2-DOF Eigenvalue Modes', 'Capsize speed (km/h), weave frequency (Hz), wobble frequency (Hz) — from linearised equations of motion'],
        ['Steering Analysis', 'Self-aligning torque (N·m), steering moment at set angle and speed'],
        ['Roll Dynamics', 'Roll angle at set speed/radius; roll rate coupling coefficient'],
        ['Wheel Spin', 'Gyroscopic precession torque from front/rear wheel spin at the set speed'],
    ],
    col_widths=[2.0, 4.3]
)

h2('Right Panel: Sweep Charts')
add_table(
    ['Chart', 'X axis', 'Y axis', 'Use'],
    [
        ['Wheel Spin vs Speed',         'Speed (km/h)', 'Gyroscopic torque (N·m)', 'Shows speed at which gyroscopic effects become significant'],
        ['Roll Angle vs Speed',         'Speed (km/h)', 'Roll angle (°)',           'How lean requirement changes with speed for fixed radius'],
        ['Steering Sweep',              'Steer angle (°)', 'Lateral force (N)',     'Steering effort vs angle for the set speed'],
        ['2-DOF Modes vs Speed',        'Speed (km/h)', 'Mode eigenvalue',         'Capsize / weave / wobble mode evolution with speed'],
    ],
    col_widths=[1.8, 1.5, 1.5, 1.5]
)

note('Capsize, weave, and wobble eigenvalues are computed from a linearised Whipple bicycle model adapted '
     'for motorcycles (Cossalter §1.4). Full Pacejka tyre model (Phase 12) will be added in a future sprint.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 16 — SWEEP COMPARE
# ══════════════════════════════════════════════════════════════════════════════
h1('17. Tab 16 — Sweep Compare')
para('The Sweep Compare tab allows up to 8 saved configurations to be plotted simultaneously on the same '
     'chart. You can sweep any parameter across its full range and compare the downstream metric for all '
     'saved configs at once. This is the most powerful tool for design-space studies.')

h2('Workflow')
for i, item in enumerate([
    'Set up your first configuration in the input tabs',
    'Click Save Config — name it (e.g. "Sport Baseline", "Race Setup", "ADV Loaded")',
    'Change parameters for a second configuration, save as another config',
    'Repeat until you have up to 8 configs saved',
    'Open Sweep Compare',
    'Select the X-axis parameter to sweep and the Y-axis metric to compare',
    'All saved configs appear as colour-coded lines on the same chart',
], 1):
    bullet(item, bold_prefix=f'Step {i}: ')

h2('Managing Saved Configs')
for item in [
    'Configs are stored in browser localStorage under mcw_saved_configs',
    'Up to 8 configs supported simultaneously',
    'Each config stores the full input state at the time of saving',
    'Delete individual configs with the X button next to each name',
    'Configs persist across page refreshes and sessions',
]:
    bullet(item)

h2('Useful Sweep Combinations')
add_table(
    ['X axis parameter', 'Y axis metric', 'Engineering use case'],
    [
        ['Rear Axle Height',  'AS%',              'Sensitivity of anti-squat to ride height setting — how much does AS% shift per mm of adjustment?'],
        ['Swingarm Length',   'Wheelbase (mm)',    'Design space: what WB results from each swingarm length option?'],
        ['Spring Rate',       'Natural Frequency', 'Spring selection guide — what rate gives the target frequency for each config mass?'],
        ['Front Sprocket',    'AS%',               'Sprocket selection optimisation across configs'],
        ['Rear Sag',          'AS%',               'How AS% changes as the bike settles under load for each config'],
        ['CoG Height',        'Wheelie g',         'Impact of luggage/fuel load on wheelie threshold'],
    ],
    col_widths=[1.7, 1.5, 3.1]
)

tip('Sweep Compare is ideal for presenting sensitivity data: one chart shows the complete design space for '
    'a parameter and lets management or the customer see the trade-off across the full range.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 17 — ANTI-DIVE
# ══════════════════════════════════════════════════════════════════════════════
h1('18. Tab 17 — Anti-Dive')
para('The Anti-Dive tab is the braking equivalent of the Anti-Squat tab. Under hard braking, inertia '
     'compresses the front fork (dive). Geometric Anti-Dive % (AD%) measures how much the fork geometry '
     'counters this. Like Anti-Squat, changes here are local overrides and do NOT modify the main store.')

h2('The Physics')
h3('Load Transfer Under Braking')
code_block('ΔW_front = M × a_brake_g × g × Y_cg / L     (weight to front, Newtons)')
para('At 1.0 g braking on a 220 kg sport bike with Y_cg = 530 mm, L = 1390 mm: '
     'ΔW ≈ 822 N — roughly doubling the front static load.')

h3('Geometric Anti-Dive %')
code_block('AD% = tan(ε) / tan(θ_front_LT) × 100')
para('Where ε = rake angle, θ_front_LT = angle of braking load transfer line from front contact patch to CoG. '
     'This is a pure geometric ratio requiring no tyre model.')

add_table(
    ['AD%', 'Behaviour', 'Typical application'],
    [
        ['0%',     'Full dive — fork compresses freely',           'Retro / cruiser bikes'],
        ['30–50%', 'Partial anti-dive — some dive, good feel',    'Road bikes, naked bikes'],
        ['50–75%', 'Sport/race geometry — noticeable anti-dive',  'Supersport, MotoGP-inspired'],
        ['> 80%',  'Aggressive — fork may jack upward under braking; reduces rider feel', 'Rarely desirable'],
    ],
    col_widths=[0.8, 2.5, 2.5]
)
note('30–60% is usually optimal. Complete anti-dive (100%) removes all fork travel feedback. Some dive is '
     'desirable because it naturally transfers rider weight forward and damps the braking load progressively.')

h2('The Eight Parameters')
add_table(
    ['#', 'Parameter', 'Direct effect on AD%', 'Note'],
    [
        ['1', 'Rear Axle Height',  'Indirect via load transfer line angle', 'Changes σ through CoG balance'],
        ['2', 'CSP Height',        'Indirect via chain reaction under braking', 'Affects rear simultaneously'],
        ['3', 'Swingarm Angle',    'Indirect via rear end behaviour during braking', 'Coupled to rear suspension response'],
        ['4', 'Swingarm Length',   'Indirect — longer arm reduces rear load transfer', 'Front sees more load duty'],
        ['5', 'Wheelbase',         'Via θ_front_LT angle (longer WB = shallower slope)', 'ΔW = M·a·h/L — longer WB reduces ΔW'],
        ['6', 'CoG Height',        'Via θ_front_LT angle steepening', 'Higher CoG → AD% changes more per g'],
        ['7', 'Rake Angle',        'DIRECT — AD% = tan(ε) / tan(θ_LT)', 'Primary control: 1° rake ≈ ±2–4% AD%'],
        ['8', 'Front Sprocket',    'Via engine braking force at combined braking', 'Relevant for linked-brake systems'],
    ],
    col_widths=[0.3, 1.5, 2.0, 2.5]
)

h2('Effective AS% at Lean Angle')
code_block('AS%_eff = AS%_upright / cos(φ)')
add_table(
    ['Lean angle φ', 'cos(φ)', 'AS%_eff if upright AS% = 80%'],
    [
        ['0° (upright)',  '1.00', '80%'],
        ['30°',          '0.866', '92%'],
        ['45°',          '0.707', '113% — jack-up risk'],
        ['50°',          '0.643', '124% — jack-up'],
    ],
    col_widths=[1.8, 1.2, 2.8]
)
warn('A bike tuned for 80% AS upright will be at 113% AS at 45° lean — above the jack-up threshold. '
     'For bikes that corner hard under power, target upright AS% of 65–75%.')

h2('IC Zone Classification')
add_table(
    ['Zone', 'IC position', 'Typical AS%', 'Application'],
    [
        ['MotoGP Zone',         'IC very high and forward', '85–100%', 'Factory GP bikes'],
        ['Sport Zone',          'IC above CoG height',      '75–90%',  'Supersport, race replica'],
        ['Naked/Standard Zone', 'IC at mid-height',         '55–75%',  'Naked, ADV, standard'],
        ['Touring Zone',        'IC low',                   '40–60%',  'Touring, cruiser'],
    ],
    col_widths=[1.8, 1.8, 1.2, 1.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 18 — SYSTEM
# ══════════════════════════════════════════════════════════════════════════════
h1('19. Tab 18 — System')
para('The System tab is a unified three-column dashboard that consolidates all parameter groups, the '
     'live 2D diagram, and a tabbed results / compare / graphs panel into one view. It is designed for '
     'experienced users who want everything visible at once.')

h2('Three-Column Layout')
add_table(
    ['Column', 'Content', 'Width'],
    [
        ['Left — Parameters', '12 collapsible parameter groups covering all geometry, suspension, chain, and mass inputs. Each slider highlights blue on hover and shows coupling — when you hover a slider, all result cards that depend on it light up.', '~35%'],
        ['Centre — Live 2D Diagram', 'Full ChassisViz2D diagram updating live as any slider moves. Shows the coupling overlay: lit results flash green when changed by the hovered parameter.', '~30%'],
        ['Right — Tabbed Results', 'Three sub-tabs: Results (all computed values), Compare (radar + bar vs another bike), Graphs (radar chart).', '~35%'],
    ],
    col_widths=[1.5, 3.5, 1.3]
)

h2('Parameter Coupling Visualisation')
para('The System tab uniquely shows which results are coupled to each parameter. When you hover over any slider:')
for item in [
    'The slider highlights blue',
    'Every result card that is mathematically dependent on that parameter lights up blue',
    'When the slider is moved, affected result cards flash green to confirm the change',
]:
    bullet(item)
para('Example: hovering over "Swingarm Length" illuminates: wheelbase output, swingarm angle, front/rear axle '
     'loads, front %, anti-squat%, IC X/Y, and chain angle — all simultaneously.')

tip('Use the System tab when you understand the coupling map and want to edit multiple parameters quickly. '
    'For beginners, the individual tabs provide more guidance and explanation per parameter.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 19 — TIRE
# ══════════════════════════════════════════════════════════════════════════════
h1('20. Tab 19 — Tire')
para('The Tire tab defines the ISO tyre dimensions and spring rates for front and rear, and computes '
     'loaded radius, deflection, contact patch length, and dynamic growth at speed.')

h2('Front Tyre Section')
add_table(
    ['Parameter', 'Range', 'Unit', 'Optimal', 'Note'],
    [
        ['Section Width',   '80–160',  'mm', '100–130', 'ISO code first number: e.g. 120/70R17 → width = 120 mm'],
        ['Aspect Ratio',    '30–90',   '%',  '55–80',   'Sidewall height as % of section width'],
        ['Rim Diameter',    '14–21',   'in', '17 in',   'ISO code last number: 17" standard sport; 21" enduro front'],
        ['Tyre Stiffness',  '80–400',  'N/mm', '150–220', 'Vertical spring rate. Soft = compliant; hard = precise but harsher'],
    ],
    col_widths=[1.7, 0.8, 0.6, 1.2, 2.0]
)
para('Live results: Free Radius (mm), Loaded Radius (mm), Deflection (mm), Contact Patch Length (mm)')
add_table(
    ['Result', 'Target range', 'Calculation'],
    [
        ['Free Radius (mm)',         '—',         'rim_radius + sidewall_height = (rim_dia × 25.4 / 2) + (width × aspect / 100)'],
        ['Loaded Radius (mm)',       '—',         'Free radius minus static deflection under axle load'],
        ['Deflection (mm)',          '< 10 mm',   'axle_load / tyre_stiffness'],
        ['Contact Patch Length (mm)', '100–145 mm front', '2 × √(2 × R_loaded × deflection) — Hertz contact approx'],
    ],
    col_widths=[2.0, 1.5, 2.8]
)

h2('Rear Tyre Section')
para('Same parameters and results as front, but typical values differ:')
add_table(
    ['Parameter', 'Typical road', 'Typical sport', 'Target contact patch'],
    [
        ['Section Width',  '150–170 mm', '190–200 mm', '130–180 mm'],
        ['Aspect Ratio',   '60–70%',     '55%',         '—'],
        ['Rim Diameter',   '17 in',      '17 in',       '—'],
        ['Tyre Stiffness', '160–200 N/mm', '190–240 N/mm', '—'],
    ],
    col_widths=[1.7, 1.5, 1.5, 1.6]
)

h2('Dynamic Growth Section')
para('At high speed, centrifugal force expands the tyre radius:')
code_block('R_dyn = R_free × (1 + k × V²)     where k = 2×10⁻⁶  [Cossalter Eq 2.4]')
para('Results shown: Dynamic Radius front/rear (mm), Combined Rate (suspension + tyre in series), '
     'and Natural Frequencies corrected with tyre compliance.')
code_block('1/k_combined = 1/k_suspension + 1/k_tyre     (series spring model)')
note('At 200 km/h (55.6 m/s): growth factor = 1 + 2e-6 × 55.6² ≈ 1.006 — a 6 mm increase in rear radius '
     'at race speeds. Significant for contact patch and trail calculations.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 20 — INERTIA
# ══════════════════════════════════════════════════════════════════════════════
h1('21. Tab 20 — Inertia')
para('The Inertia tab displays the three principal moments of inertia computed from the mass component '
     'list using the parallel-axis theorem (point-mass approximation per Cossalter Ch. 1). '
     'No inputs are available here — edit mass components in the Mass tab.')

h2('Three Inertia Axes')
add_table(
    ['Axis', 'Symbol', 'Physical meaning', 'Governs'],
    [
        ['Pitch', 'I_yy (I_pitch)', 'About lateral Y-axis (nose-up / nose-down)', 'Acceleration squat and braking dive tendency'],
        ['Roll',  'I_xx (I_roll)',  'About longitudinal X-axis (side-to-side)',    'Cornering agility, lean-in speed, weave mode'],
        ['Yaw',   'I_zz (I_yaw)',  'About vertical Z-axis (direction change)',     'Weave mode, direction-change agility'],
    ],
    col_widths=[0.8, 1.7, 2.5, 1.8]
)

h2('Formulas')
code_block('I_pitch = Σ mᵢ × [(xᵢ − X̄)² + (yᵢ − Ȳ)²] / 10⁶    [kg·m²]')
code_block('I_roll  = Σ mᵢ × [(yᵢ − Ȳ)² + (zᵢ − Z̄)²] / 10⁶    (zᵢ = 0, symmetric)')
code_block('I_yaw   = Σ mᵢ × [(xᵢ − X̄)² + (zᵢ − Z̄)²] / 10⁶')
code_block('Radius of gyration k = √(I / M_total)    [m displayed as mm]')

h2('Reading the Display')
para('Three colour-coded cards (green / amber / red based on relative magnitude to the largest axis):')
for item in [
    'Large absolute value (red card) = dominant inertia axis; hardest to rotate about this axis',
    'Bar chart below shows relative magnitude: longest bar = dominant, shortest = most agile axis',
    'For a motorcycle, I_pitch is typically largest (long wheelbase, distributed mass fore-aft)',
    'I_yaw governs weave mode — lower I_yaw = more agile direction change',
]:
    bullet(item)

h2('Engineering Significance')
add_table(
    ['Inertia', 'Effect when HIGH', 'Effect when LOW'],
    [
        ['I_pitch', 'Slower pitch response under braking/accel; more stable on bumps', 'Quick pitch; susceptible to wheelies'],
        ['I_roll',  'Slower to lean; more stable at high speed in corners',             'Quick lean; agile cornering'],
        ['I_yaw',   'More directional stability; harder to change direction',            'Agile; susceptible to weave mode'],
    ],
    col_widths=[1.2, 2.9, 2.9]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 21 — STABILITY
# ══════════════════════════════════════════════════════════════════════════════
h1('22. Tab 21 — Stability')
para('The Stability tab computes and displays wheelie/stoppie thresholds, lean clearance, minimum turning '
     'radius, maximum climbable grade, and the full set of handling indices. Three adjustable inputs '
     'control the analysis conditions.')

h2('Section: Wheelie / Stoppie Thresholds')
code_block('a_wheelie = g × (WB − X_cg) / Y_cg    [m/s² or g]     [Foale Ch.5]')
code_block('a_stoppie = g × X_cg / Y_cg            [m/s² or g]')
add_table(
    ['Threshold', 'Target', 'If too low', 'If too high'],
    [
        ['Wheelie (g)', '> 0.8 g preferred', 'Front lifts easily — stability risk; reduces braking effectiveness', 'Hard to perform power wheelies; typically not an issue for road bikes'],
        ['Stoppie (g)', '> 1.0 g preferred', 'Rear lifts under moderate braking — ABS may not prevent', 'Good — means stable under hard braking'],
    ],
    col_widths=[1.5, 1.5, 2.1, 2.1]
)
note('Both thresholds are purely geometric — they depend on CoG position and wheelbase. '
     'Edit CoG height and X position in the Mass tab to change them.')

h2('Section: Lean & Turning')
add_table(
    ['Parameter / Result', 'Description', 'Target'],
    [
        ['Footpeg Lateral Offset (mm)', 'Input: distance from centreline to footpeg tip. Sets lean clearance.', '300–400 mm'],
        ['Friction Coeff μ', 'Input: tyre-road friction. 0.8 dry; 0.5 wet; 1.0 slick.', '0.75–0.95 dry'],
        ['Steering Lock Angle (°)', 'Input: max steering angle for turning radius calculation.', '30–42°'],
        ['Lean Limit (°)', 'Result: max lean before footpeg grounds = arctan(clearance / offset)', '44–56° road'],
        ['Min Turn Radius (m)', 'Result: WB / sin(lock_angle)', '< 5 m ideal'],
        ['Turning Circle Ø (m)', 'Result: 2 × Min Turn Radius', '—'],
        ['Max Grade (°)', 'Result: arctan(μ) — steepest slope tyre can climb', 'arctan(0.8) = 38.7°'],
    ],
    col_widths=[2.0, 2.8, 1.5]
)

h2('Section: Handling Indices')
add_table(
    ['Index', 'Formula', 'Target', 'Interpretation'],
    [
        ['Stability Index (SI)', 'trail × WB / 10⁶', '0.08–0.20', 'Higher = more straight-line stable'],
        ['Agility Index (AI)',   'I_yaw / (M × WB²)', '< 0.15',   'Lower = more agile, quicker turns'],
        ['Wobble Sensitivity',   '10⁶ / (trail × WB)', '< 15',     'Lower = less prone to wobble'],
        ['Pitch Sensitivity',    'X_cg / WB² (%/mm)',  '—',        'Weight split change per mm WB change'],
        ['Rear Squat (mm)',      'travel × (1 − AS%/100)', '< 20 mm ideal', 'Actual rear compression under current accel'],
        ['Fork Dive (mm)',       'travel × (1 − AD%/100)', '< 30 mm ideal', 'Actual fork compression under current braking'],
    ],
    col_widths=[1.7, 2.0, 1.1, 1.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 22 — FORK COMPLIANCE
# ══════════════════════════════════════════════════════════════════════════════
h1('23. Tab 22 — Fork Compliance')
para('The Fork Compliance tab analyses fork structural flexibility and its effect on steering geometry '
     'under braking loads. Flexible forks cause the front axle to deflect laterally and twist the '
     'steering system — both effects change effective trail and introduce self-steer torque.')

h2('Section: Fork Stiffness Inputs')
add_table(
    ['Parameter', 'Range', 'Unit', 'Optimal', 'Note'],
    [
        ['Fork Bending Stiffness', '10–120', 'N/mm', '35–65', 'Lateral stiffness at the axle. Sport: 40–60; MX: 25–40; enduro: 20–35'],
        ['Fork Torsional Stiffness', '100–1500', 'N·m/°', '350–650', 'Twist resistance. Higher = sharper steering feel; lower = flex under cornering load'],
        ['Steering Head Stiffness', '200–3000', 'N·m/°', '600–1500', 'Frame compliance at the head tube. Very stiff frames > 2000 N·m/°'],
    ],
    col_widths=[1.9, 0.9, 0.9, 1.0, 1.6]
)

h2('Section: Braking Analysis Results')
para('All results are computed at 0.8 g deceleration, 70% front brake share, 5° reference steer angle:')
add_table(
    ['Result', 'Formula', 'Target / Warning level'],
    [
        ['Front Brake Force (N)', 'M × 0.8g × 0.70', '—'],
        ['Fork Deflection (mm)', 'F_brake / k_bend', '< 2.5 mm ideal'],
        ['Steering Torque (N·m)', 'R_f × (trail/1000) × sin(δ)', '—'],
        ['Trail Change Δ (mm)', '−deflection × cos(rake)', 'Reduces effective trail under braking'],
        ['Effective Trail (mm)', 'Trail + Δ', 'Should stay within 60–140 mm'],
        ['Steer Flex Angle (°)', 'torque / k_torsional', '< 0.5° ideal; > 1.5° dangerous'],
    ],
    col_widths=[1.9, 2.0, 2.4]
)

h2('Classification Display')
para('Three possible messages (coloured box):')
for item, color in [
    ('Green (✓): Steer flex < 0.5° — within acceptable limits', 'Green'),
    ('Amber (⚠): Steer flex 0.5–1.5° — perceptible by rider', 'Amber'),
    ('Red (✕): Steer flex > 1.5° — potentially dangerous [Cossalter Ch.6]', 'Red'),
]:
    bullet(item)
warn('A red classification means the fork flexibility is large enough to cause unintended self-steer '
     'under hard braking. Increase fork OD, wall thickness, or select a stiffer material in the FEM tab.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 23 — AERO
# ══════════════════════════════════════════════════════════════════════════════
h1('24. Tab 23 — Aero (Aerodynamics)')
para('The Aero tab is a dedicated full-width aerodynamics workbench. It exposes more detail than the '
     'aerodynamics section in the Dynamics tab — three-column layout with inputs, single-point KPIs, '
     'and four sweep charts vs speed.')

h2('Left Column: Inputs')
h3('Aero Coefficients section')
add_table(
    ['Parameter', 'Range', 'Typical values'],
    [
        ['Drag Coeff Cx', '0.15–1.0', 'Faired sport: 0.33–0.40; naked: 0.50–0.65; scooter: 0.75–0.90'],
        ['Lift Coeff Cz', '−0.5 to +0.5', 'Racing fairing: −0.05 to −0.15 (downforce); touring: 0 to +0.10'],
        ['Frontal Area (m²)', '0.15–1.0', 'Sport: ~0.33; naked: ~0.50; ADV: ~0.65; scooter: ~0.70'],
        ['Pressure Centre X (mm)', '300–1200', 'From front axle. If forward of CoG → nose-down pitch moment under aero load'],
    ],
    col_widths=[1.9, 1.4, 3.0]
)

h3('Engine & Speed section')
add_table(
    ['Parameter', 'Range', 'Note'],
    [
        ['Engine Power (kW)', '5–300', 'Peak shaft power for top speed prediction'],
        ['Drivetrain η',      '0.70–0.98', 'Power delivery efficiency. Chain loss ≈ 10–12%'],
        ['Reference Speed (km/h)', '50–350', 'Speed used for single-point KPI column output'],
        ['Max Chart Speed (km/h)', '100–400', 'Upper limit of X axis on all four sweep charts'],
    ],
    col_widths=[2.0, 1.2, 3.0]
)

h2('Centre Column: Single-Point KPIs')
para('All computed at the Reference Speed:')
for item in [
    'Drag Force (N)',
    'Lift Force (N) — negative = downforce',
    'Drag Power (kW)',
    'Aero Pitch Moment (N·m)',
    'ΔW Front aero (N) — aero load transfer to/from front axle',
    'Dynamic Pressure (Pa) = ½ρV²',
    'Top Speed prediction (km/h and m/s)',
    'Drag at 100 km/h (N) — for regulatory / efficiency comparisons',
]:
    bullet(item)
code_block('V_max = ∛(2 × P × η / (ρ × Cx × A))     [Cossalter Ch.4/8]')

h2('Right Column: Four Sweep Charts')
add_table(
    ['Chart', 'Y axis', 'Reference line'],
    [
        ['Drag Force vs Speed', 'Drag (N)',            'Dashed at reference speed'],
        ['Lift Force vs Speed', 'Lift/Downforce (N)',  'Dashed at reference speed'],
        ['Drag Power vs Speed', 'Power (kW)',          'Dashed at reference speed'],
        ['Front Aero Load Transfer vs Speed', 'ΔW Front (N)', 'Dashed at reference speed'],
    ],
    col_widths=[2.2, 1.8, 2.3]
)
note('The aero formulas at the bottom of the panel: F_D = ½ρCxAV²  ·  F_L = ½ρCzAV²  ·  '
     'P = F_D×V  ·  M_aero = F_L×(x_cp − X_cg)  ·  ΔW_front = M_aero / WB')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SIGN CONVENTIONS
# ══════════════════════════════════════════════════════════════════════════════
h1('25. Sign Conventions & Coordinate Systems')

h2('Swingarm Angle — CW+ Convention')
para('All displayed swingarm angles in the workbench use the CW+ (Clockwise Positive) convention:')
for item in [
    'Positive (+) angle = swingarm slopes downward toward the rear axle = CW rotation in side view',
    'Typical sport bike: +4° to +8° displayed',
    'Negative angle would mean rear axle ABOVE swingarm pivot — only in unusual off-road configurations',
]:
    bullet(item)
note('Internally the physics engines use an ACW+ convention (arcsin result). The display negates the raw value '
     'at all display sites. This is mentioned in all three sections: Geometry results bar, Anti-Squat tab, '
     'Anti-Dive tab.')

h2('Coordinate Systems by Module')
add_table(
    ['Module', 'Origin', '+X direction', '+Y direction', 'Notes'],
    [
        ['ChassisViz2D (2D drawing)', 'Swingarm pivot', 'Forward (toward front)', 'Upward', 'Screen: +X right, +Y down. Physics +X maps to screen LEFT (mirror).'],
        ['antiSquat.ts (IC computation)', 'Front tyre contact patch (0,0)', 'Rearward', 'Upward', 'Rear contact patch = (WB, 0)'],
        ['computeAll.ts / CoG outputs', 'Front axle', 'Rearward', 'Upward', 'X_cg = distance from front axle rearward; Y_cg = height from ground'],
    ],
    col_widths=[1.8, 1.5, 1.2, 1.2, 1.6]
)

h2('2D Diagram Orientation')
para('Standard motorcycle engineering side-view convention:')
for item in [
    'Front wheel: LEFT side of diagram',
    'Rear wheel: RIGHT side of diagram',
    'Ground: horizontal line at the bottom',
    'Viewing from the RIGHT side of the bike, looking LEFT',
]:
    bullet(item)

h2('Chain Convention')
add_table(
    ['Chain run', 'Name in code', 'Colour in diagram', 'Physical meaning'],
    [
        ['Upper run', 'CHAIN_TENSION (A→B)', 'Thick cyan', 'Tension run — carries the driving force; used for IC construction'],
        ['Lower run', 'CHAIN_SLACK (A→B)', 'Thin grey', 'Slack run — returns chain to countershaft'],
    ],
    col_widths=[1.2, 2.0, 1.4, 2.2]
)

h2('Positive Angle Conventions Summary')
add_table(
    ['Angle', 'Positive direction', 'Typical value'],
    [
        ['Rake (ε)',            'Leaning back from vertical (clockwise)',   '24–30°'],
        ['Swingarm angle',      'Downward slope from pivot to axle (CW)',  '+4° to +8°'],
        ['Chain force angle',   'Upward from CS toward rear (CCW)',        'Positive for most bikes'],
        ['Bank angle (φ)',      'Lean away from vertical (either side)',    '0–55°'],
        ['Steer angle (δ)',     'Turn toward right (CW viewed from top)',   'Positive right'],
    ],
    col_widths=[1.8, 2.8, 1.7]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# QUICK REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
h1('26. Quick Reference — Parameter Sensitivity Tables')

h2('Anti-Squat Sensitivity (approximate, single-parameter changes)')
add_table(
    ['Change', 'AS% effect', 'Direction'],
    [
        ['+10 mm CSP height (countershaft)',   '+4 to +7%',  '↑'],
        ['+1° swingarm angle (CW steeper)',     '+4 to +6%',  '↑'],
        ['+1 tooth front sprocket',             '+3 to +5%',  '↑'],
        ['+1 tooth rear sprocket',              '−2 to −4%',  '↓'],
        ['+10 mm rear axle height',             '+3 to +5%',  '↑'],
        ['+40 mm swingarm length',              '−3 to −5%',  '↓'],
        ['+100 mm wheelbase',                   '−2 to −4%',  '↓'],
        ['+50 mm CoG height',                   '−2 to −4%',  '↓'],
    ],
    col_widths=[3.5, 1.5, 0.7]
)

h2('Steering Geometry Sensitivity')
add_table(
    ['Change', 'Trail effect'],
    [
        ['+1° rake (more reclined)', '+4 to +6 mm trail'],
        ['+5 mm fork offset',        '−5 to −8 mm trail'],
        ['+10 mm front wheel radius','+6 to +10 mm trail'],
        ['+1° rake',                 'AD% +2 to +4%'],
    ],
    col_widths=[3.5, 2.8]
)

h2('Suspension Sensitivity')
add_table(
    ['Change', 'Natural frequency effect'],
    [
        ['+10% spring rate',     '+4.9% natural frequency'],
        ['+10% sprung mass',     '−5.1% natural frequency'],
        ['MR 0.9 → 0.8',        'WR drops 20%, frequency drops 10.6%'],
        ['+25% damping coeff',   'ζ increases 25%, no frequency change'],
    ],
    col_widths=[3.5, 2.8]
)

h2('Anti-Dive Sensitivity')
add_table(
    ['Change', 'AD% effect'],
    [
        ['+1° rake angle',   '+2 to +4% AD%'],
        ['+100 mm wheelbase','−1 to −3% AD% (via θ_LT angle)'],
        ['+50 mm CoG height','Variable — steepens θ_LT, changes AD%'],
    ],
    col_widths=[3.5, 2.8]
)

h2('Target Ranges — Quick Lookup')
add_table(
    ['Metric', 'Good', 'Acceptable', 'Action required'],
    [
        ['Trail (mm)',              '80–120',     '60–150',      '< 60 or > 150'],
        ['Anti-Squat%',            '80–120%',    '60–140%',     '< 40% or > 150%'],
        ['Anti-Dive%',             '30–60%',     '20–75%',      '< 15% or > 85%'],
        ['Front weight %',         '48–55%',     '42–60%',      '< 38% or > 65%'],
        ['Front natural freq (Hz)','1.3–1.7',    '0.9–2.0',     '< 0.7 or > 2.2'],
        ['Front sag %',            '25–30%',     '20–35%',      '< 15% or > 40%'],
        ['Lean limit (°)',         '44–56°',     '38–62°',      '< 35°'],
        ['Frame safety factor',    '≥ 3.0',      '1.5–3.0',     '< 1.5'],
        ['Fork steer flex (°)',    '< 0.5°',     '0.5–1.5°',    '> 1.5°'],
        ['Knee angle (°)',         '90–130°',    '70–150°',     '< 70° or > 155°'],
    ],
    col_widths=[2.2, 1.3, 1.5, 1.3]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# EXPORT & PERSISTENCE
# ══════════════════════════════════════════════════════════════════════════════
h1('27. Export & Data Persistence')

h2('Session Auto-Save')
para('The current input state is automatically saved to browser localStorage under the key mcw_session '
     'whenever any parameter changes. When you reload the app, it restores the last session automatically. '
     'This includes the family name and all parameter values.')

h2('JSON Export')
para('Exports the complete ComputeAllInput + ComputeAllResult as a structured JSON file. Contains:')
for item in [
    'All 23-tab input parameters (geometry, mass components, suspension, chain, ergo, dynamics, tire, FEM, aero, etc.)',
    'All computed results from all 13 physics engines',
    'Timestamp and family name',
]:
    bullet(item)

h2('CSV Export')
para('Exports a flattened key-value table of all parameters and results. Suitable for:')
for item in [
    'Import into Microsoft Excel for further analysis',
    'Python/pandas data processing',
    'Database archiving of design iterations',
]:
    bullet(item)

h2('Sweep CSV Export (Chassis Sim tab)')
para('Exports the full suspension sweep dataset (motion ratio, wheel rate, AS%, trail at every travel mm) '
     'and the dynamics sweep dataset as separate CSV downloads.')

h2('Saved Configs (Sweep Compare)')
add_table(
    ['Storage key', 'Contents', 'Max count', 'When cleared'],
    [
        ['mcw_saved_configs', 'SavedConfig[] — name + full input snapshot', '8', 'Manually via delete button or clear localStorage'],
        ['mcw_custom_bikes',  'CustomBike[] — name + full input snapshot',  'Unlimited', 'Manually via delete in the modal'],
        ['mcw_session',       'Current { input, familyName } state',         '1', 'Overwritten on every change'],
    ],
    col_widths=[1.8, 2.5, 1.0, 1.9]
)

h2('Custom Bikes')
para('The Custom Bike modal (gear icon in header) lets you:')
for item in [
    'Save the current configuration as a named custom bike',
    'Rename an existing custom bike',
    'Overwrite a custom bike with the current configuration',
    'Delete a custom bike from the library',
    'Load a custom bike — replaces the entire input state',
]:
    bullet(item)
para('Custom bikes appear in the Family dropdown alongside the 8 presets and in the Compare tab Bike B selector.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# BIKE PRESETS
# ══════════════════════════════════════════════════════════════════════════════
h1('28. Bike Presets (8 Families)')
para('The 8 preset families in families.ts represent validated starting points for each major motorcycle '
     'category. All values were corrected on 2026-04-08.')

add_table(
    ['Family', 'Rake', 'Fork Offset', 'WB', 'SA Length', 'Front Travel', 'Rear Travel', 'Front Sprocket', 'Rear Sprocket', 'Notes'],
    [
        ['Sport / Supersport', '24°', '35 mm', '1390 mm', '490 mm', '320 mm', '220 mm', '14T', '42T', 'Clip-on bars, 95 mm trail'],
        ['Naked / Roadster',   '25°', '38 mm', '1430 mm', '495 mm', '120 mm', '130 mm', '15T', '42T', 'Standard bars, 100 mm trail'],
        ['Adventure / ADV',    '27°', '44 mm', '1500 mm', '520 mm', '200 mm', '220 mm', '15T', '42T', 'Riser bars, 120 mm trail'],
        ['Cruiser',            '30°', '35 mm', '1570 mm', '655 mm', '130 mm', '120 mm', '15T', '45T', 'Pull-back bars, 115 mm trail, long SA'],
        ['Touring / Luxury',   '30°', '38 mm', '1620 mm', '692 mm', '120 mm', '110 mm', '15T', '43T', 'Pull-back bars, 120 mm trail'],
        ['Supermoto',          '25°', '40 mm', '1390 mm', '470 mm', '270 mm', '180 mm', '13T', '42T', 'Wide bars, 105 mm trail, lightweight'],
        ['Enduro / Off-Road',  '26°', '32 mm', '1460 mm', '495 mm', '300 mm', '280 mm', '13T', '52T', '21" front, 120 mm trail'],
        ['Scooter / Urban',    '26°', '44 mm', '1360 mm', '350 mm', '90 mm',  '100 mm', 'CVT', 'CVT', '90 mm trail, CVT mode on'],
    ],
    col_widths=[1.5, 0.5, 0.8, 0.8, 0.8, 0.9, 0.9, 0.9, 0.9, 0.9]
)

h2('Preset Correction History (2026-04-08)')
for item in [
    'Scooter: forkOffset corrected 75 mm → 44 mm; rearSprocket corrected 68T → CVT',
    'Cruiser: swingarmLength corrected 640 mm → 655 mm',
    'Touring: swingarmLength corrected 650 mm → 692 mm',
    'Enduro: forkOffset corrected 22 mm → 32 mm',
    'ALL 8 families: rider mass Y position += 200 mm (CoG at hip, not seat surface)',
]:
    bullet(item)

h2('Physics Reference Documents')
para('All formulas in MPAW are traceable to:')
for item in [
    'Foale, T. — "Motorcycle Handling and Chassis Design" — Ch. 5 (IC construction, trail), Ch. 11 (stability)',
    'Cossalter, V. — "Motorcycle Dynamics" 2006 — Ch. 1–8 (squat ratio R, eigenvalue stability, tyre dynamics, aero)',
    'anti_squat_anti_dive_motorcycles.docx.pdf — §1–§9 (load transfer, AD% geometric formula, lean analysis)',
]:
    bullet(item)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
output_path = '/home/dikshant/Desktop/Moter_bike/Chassis_Workbench_User_Guide.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
