"""
Generate a DOCX report of everything built in the Moter_bike project.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)

# ── Helper functions ──────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x77, 0x7C)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11) if p.runs else None
    return p

def bullet(text, bold_part=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_part and text.startswith(bold_part):
        run1 = p.add_run(bold_part)
        run1.bold = True
        run1.font.size = Pt(11)
        rest = text[len(bold_part):]
        if rest:
            run2 = p.add_run(rest)
            run2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x17, 0x17, 0x17)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(shading)
    return p

def divider():
    doc.add_paragraph("─" * 70)

def spacer():
    doc.add_paragraph("")

# ═══════════════════════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════════════════════
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run("Motorcycle Chassis Workbench")
run.font.size  = Pt(26)
run.font.bold  = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Complete Project Report — What We Built & How It Works")
r.font.size  = Pt(14)
r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr = date_p.add_run(f"Generated: {datetime.date.today().strftime('%B %d, %Y')}")
dr.font.size  = Pt(11)
dr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
dr.font.italic = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
h1("1. What Is This Project?")
body(
    "The Motorcycle Chassis Workbench is a complete software toolkit for designing, "
    "simulating, and analysing motorcycle geometry and physics. It has two main parts:"
)
bullet("A web application (runs in your browser) where you can visually adjust motorcycle parts and instantly see how physics changes.", "A web application")
bullet("A Python physics engine that does all the maths — calculating things like trail, anti-squat, centre of gravity, and more.", "A Python physics engine")
spacer()
body(
    "Think of it like a digital motorcycle workshop. You change one thing (for example, "
    "the swingarm length) and the entire bike recalculates automatically — wheelbase changes, "
    "weight distribution shifts, anti-squat percentage updates. Everything is connected."
)

spacer()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. THE WEB APP
# ═══════════════════════════════════════════════════════════════════════════════
h1("2. The Web Application (chassis-workbench)")
body(
    "The web app is built with TypeScript, React, and Vite. It runs on your computer "
    "using Electron (so it looks and feels like a desktop app). It has 17 tabs, each "
    "covering a different aspect of the motorcycle."
)

h2("2.1  The 17 Tabs")
tabs = [
    ("Geometry",         "Shows the basic shape — rake angle, trail, wheelbase, fork offset."),
    ("Anti-Squat",       "Calculates how much the bike squats under acceleration. Shows the Instant Centre point."),
    ("Anti-Dive",        "Calculates how much the front dips under braking."),
    ("Centre of Gravity","Shows where the combined weight of bike + rider sits."),
    ("Weight Transfer",  "Shows how weight shifts front-to-back when accelerating or braking."),
    ("Chain & Sprockets","Models the drive chain geometry and forces."),
    ("Suspension Travel","Sweeps through suspension movement and shows motion ratio."),
    ("Four-Bar Linkage", "Models complex suspension linkages used on some bikes."),
    ("Stability",        "Linearised stability analysis — shows weave, wobble, capsize modes."),
    ("Aerodynamics",     "Drag force, lift, pitch moment, top speed estimates."),
    ("Tyre (Pacejka)",   "Magic Formula tyre model for grip and slip."),
    ("MBD 2D",           "Multi-Body Dynamics simulation in 2 dimensions."),
    ("MBD 3D",           "Multi-Body Dynamics simulation in 3 dimensions."),
    ("Sweep",            "Sweep any parameter across a range and plot how outputs change."),
    ("Family Presets",   "Quick-load Sport / Naked / Cruiser motorcycle configurations."),
    ("Graphs",           "Plot any computed value over time or across sweep ranges."),
    ("Report",           "Export all results to a printable summary."),
]
for name, desc in tabs:
    bullet(f"{name}: {desc}", f"{name}:")

spacer()

h2("2.2  How the Web App Works Internally")
body(
    "The app uses a state management system (Zustand) where every parameter is stored "
    "centrally. When you change a value in any tab, the physics engine recalculates "
    "everything instantly and all tabs update at the same time. The physics formulas "
    "live in TypeScript files inside the src/engine/ folder."
)

spacer()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. PYTHON PHYSICS ENGINE
# ═══════════════════════════════════════════════════════════════════════════════
h1("3. The Python Physics Engine (dynamics_engine)")
body(
    "We built a standalone Python package that does the same physics calculations as "
    "the web app, but in pure Python. This is useful for running batch calculations, "
    "parameter sweeps, and automated testing without opening a browser."
)

h2("3.1  What Makes It Special — The DAG System")
body(
    "DAG stands for Directed Acyclic Graph. This is the clever part. Instead of writing "
    "long chains of calculations manually, every output is declared as a node that "
    "depends on other nodes. When you change one input, the engine automatically figures "
    "out the correct order to recalculate everything else."
)
body("Example chain:")
code_block(
    "swingarm_length\n"
    "   → wheelbase\n"
    "       → x_cg  (centre of gravity X position)\n"
    "           → front_weight_pct\n"
    "               → anti_squat_pct\n"
    "                   → wheelie_threshold_g"
)
body(
    "Change swingarm_length and all 5 downstream values update automatically "
    "in the correct order."
)

spacer()

h2("3.2  The Physics Formulas We Implemented")
body("These match exactly what the TypeScript web app uses:")

formulas = [
    ("Trail",
     "Trail = (R_front × sin(rake) − fork_offset) / cos(rake)\n"
     "rake is measured from vertical. A sport bike typically gives ~97 mm trail."),
    ("Mechanical Trail",
     "Mechanical Trail = Trail / cos(rake)\n"
     "This is the effective lever arm for steering feel."),
    ("Swingarm Angle",
     "Angle = arcsin( (rear_axle_height − pivot_height) / swingarm_length )\n"
     "Uses arcsin because swingarm_length is the hypotenuse, not the base."),
    ("Anti-Squat Instant Centre",
     "Two lines are drawn: one along the swingarm, one along the chain force direction.\n"
     "Where they cross is the Instant Centre (IC). The maths uses slope-form line intersection."),
    ("Anti-Squat Percentage",
     "A line is drawn from the IC to the rear contact patch.\n"
     "AS% = height of that line at the front wheel ÷ CoG height × 100\n"
     "100% means the chain force perfectly counteracts squat."),
    ("Centre of Gravity",
     "X_cg = sum(mass × x_position) / total_mass\n"
     "Y_cg = sum(mass × y_position) / total_mass\n"
     "Each component (engine, frame, rider, fuel, etc.) contributes."),
    ("Wheelie Threshold",
     "The acceleration (in g) at which the front wheel lifts off.\n"
     "wheelie_g = (x_cg / wheelbase) × (1 / Y_cg × wheelbase × 9.81)"),
]

for name, explanation in formulas:
    h3(f"  {name}")
    code_block(explanation)

spacer()

h2("3.3  The Three Motorcycle Presets")
body(
    "We created three presets that mirror the TypeScript families.ts exactly, "
    "so Python and browser results always match:"
)
presets = [
    ("Sport",   "Aggressive geometry. Rake 24°, short wheelbase ~1415 mm, low CoG. "
                "Like a Yamaha R1 or Honda CBR."),
    ("Naked",   "Upright riding position. Rake 26°, medium wheelbase ~1450 mm. "
                "Like a Kawasaki Z900."),
    ("Cruiser", "Relaxed, long geometry. Rake 32°, long wheelbase ~1600 mm. "
                "Like a Harley-Davidson."),
]
for name, desc in presets:
    bullet(f"{name}: {desc}", f"{name}:")

spacer()

h2("3.4  The Sweep Tool")
body(
    "You can tell the engine to vary one input across a range and record how "
    "all outputs change. For example, sweeping swingarm_length from 480 mm to 700 mm "
    "produces a 7-panel chart showing how wheelbase, CoG, weight distribution, "
    "anti-squat, wheelie threshold, trail, and stability all change together."
)
code_block(
    "sweep = model.sweep('swingarm_length', 480, 700, steps=30,\n"
    "                    output_names=['wheelbase', 'front_pct',\n"
    "                                  'anti_squat_pct', 'trail'])"
)

spacer()

h2("3.5  Output Files")
outputs = [
    ("dynamics_engine/motorcycle_dynamics.py",    "Main engine — ~1700 lines, all physics"),
    ("dynamics_engine/__init__.py",               "Package entry point"),
    ("dynamics_engine/sweep_swingarm_length.png", "7-panel sweep chart"),
    ("dynamics_engine/parameter_dag.png",         "Visual diagram of the DAG network"),
]
for fname, desc in outputs:
    bullet(f"{fname}  —  {desc}", fname)

spacer()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. MBD ENGINE
# ═══════════════════════════════════════════════════════════════════════════════
h1("4. Multi-Body Dynamics Engine (mbd_engine)")
body(
    "Multi-Body Dynamics (MBD) is advanced simulation where the motorcycle is modelled "
    "as a collection of rigid bodies (frame, swingarm, wheels, rider) connected by joints "
    "(revolute, prismatic, spherical). The system solves Newton's laws for all bodies "
    "simultaneously at each time step."
)

h2("4.1  What We Built")
bullet("2D MBD Solver: simulates motion in the vertical plane (pitch, heave, suspension travel).", "2D MBD Solver:")
bullet("3D MBD Solver: full 6-degree-of-freedom simulation including roll, yaw, and lean.", "3D MBD Solver:")
bullet("Generalized-α Solver: a professional-grade time integrator used in engineering software.", "Generalized-α Solver:")
bullet("Newton-Raphson Solver: iterative solver for the nonlinear constraint equations.", "Newton-Raphson Solver:")
bullet("Assembler 2D & 3D: builds the mass matrix, Jacobian, and force vector each time step.", "Assembler 2D & 3D:")

h2("4.2  Test Coverage")
body(
    "32 automated tests were written across 3 phases and all pass:"
)
phases = [
    ("Phase 1", "Basic rigid body motion, constraint enforcement, energy conservation."),
    ("Phase 2", "Multi-body assemblies, joint reactions, constraint Jacobian accuracy."),
    ("Phase 3", "3D quaternion dynamics, angular momentum, numerical stability."),
]
for phase, desc in phases:
    bullet(f"{phase}: {desc}", f"{phase}:")

spacer()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. CHASSIS SIMULATION ENGINE
# ═══════════════════════════════════════════════════════════════════════════════
h1("5. Chassis Simulation Engine (chassis_sim)")
body(
    "A higher-level simulation library sitting above the raw MBD engine. "
    "It speaks 'motorcycle language' — you give it geometry parameters and "
    "it gives back suspension curves, dynamics, and stability results."
)

h2("5.1  Modules")
modules = [
    ("geometry.py",   "Computes wheelbase, trail, CoG, chain geometry from basic dimensions."),
    ("dynamics.py",   "Weight transfer under braking and acceleration, anti-dive, anti-squat sweep."),
    ("sweep.py",      "Suspension travel sweep — motion ratio, wheel rate, anti-squat vs travel."),
    ("stability.py",  "Linearised Meijaard bicycle model — computes weave, wobble, capsize eigenvalues."),
    ("fourbar.py",    "Four-bar linkage kinematics for progressive suspension."),
]
for name, desc in modules:
    bullet(f"{name}: {desc}", f"{name}:")

h2("5.2  REST API (FastAPI)")
body(
    "All chassis_sim functions are exposed as HTTP endpoints so the web app "
    "can call them. The API runs locally on your machine."
)
endpoints = [
    ("POST /api/geometry",     "Calculate all geometry values from input dimensions."),
    ("POST /api/anti-squat",   "Calculate anti-squat percentage and instant centre."),
    ("POST /api/dynamics",     "Run braking/acceleration dynamics sweep."),
    ("POST /api/sweep",        "Run suspension travel sweep."),
    ("POST /api/stability",    "Run stability analysis, return eigenvalues and mode shapes."),
]
for endpoint, desc in endpoints:
    bullet(f"{endpoint}  —  {desc}", endpoint)

spacer()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. KNOWLEDGE GRAPH (MIND MAP)
# ═══════════════════════════════════════════════════════════════════════════════
h1("6. Project Knowledge Graph — The Mind Map")
body(
    "We used a tool called graphify to read every source file in the project and "
    "automatically build a knowledge graph — a map of how everything connects to "
    "everything else. Think of it as an X-ray of the codebase."
)

h2("6.1  What the Graph Contains")
bullet("1,194 nodes — each node is a function, class, file, or concept.", "1,194 nodes")
bullet("2,537 edges — each edge is a relationship (calls, imports, contains, uses).", "2,537 edges")
bullet("77 communities — groups of closely related nodes (like brain regions).", "77 communities")
bullet("124 source files scanned — TypeScript, Python, React, and test files.", "124 source files scanned")

h2("6.2  How to Use It")
body("Two commands:")
code_block(
    "# Open the interactive brain-style mind map in your browser\n"
    "brainmap\n\n"
    "# Ask the graph a question\n"
    'python3.13 -m graphify query "how does anti-squat work"\n\n'
    "# Find the path between two things\n"
    'python3.13 -m graphify path "MassComponent" "anti_squat_pct"'
)

h2("6.3  The Interactive Brain Map (mindmap.html)")
body(
    "We built a custom animated visualisation that looks like a neural network / brain:"
)
features = [
    "Nodes glow in colour — each colour is a different community (module group).",
    "Neural signal particles stream along the edges continuously, like electrical impulses.",
    "Click any node — dims everything else, highlights its connections, fires signals outward.",
    "Right panel shows: file name, location, type, number of connections, community.",
    "Connections panel lists all neighbours with arrows showing direction — click to jump.",
    "Search box — type any function or class name, click result to fly to it.",
    "Community legend — click to hide or show entire groups.",
    "Zoom with scroll, pan by dragging, ⚡ Burst fires 40 signals at once.",
    "Never stops — gentle breathing animation keeps the graph alive forever.",
]
for f in features:
    bullet(f)

spacer()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. FILE STRUCTURE
# ═══════════════════════════════════════════════════════════════════════════════
h1("7. Project File Structure")
body("Here is what lives where on your computer:")
code_block(
    "Moter_bike/\n"
    "├── chassis-workbench/       ← Web app (TypeScript + React + Electron)\n"
    "│   ├── src/\n"
    "│   │   ├── engine/          ← Physics formulas in TypeScript\n"
    "│   │   ├── components/      ← React UI — all 17 tabs\n"
    "│   │   ├── store/           ← State management (Zustand)\n"
    "│   │   └── data/            ← Presets (Sport, Naked, Cruiser)\n"
    "│   └── CLAUDE.md            ← Architecture documentation\n"
    "│\n"
    "├── dynamics_engine/         ← Python DAG physics engine\n"
    "│   ├── motorcycle_dynamics.py  ← All physics, ~1700 lines\n"
    "│   └── __init__.py\n"
    "│\n"
    "├── mbd_engine/              ← Multi-body dynamics solver\n"
    "│   ├── core/                ← Assembler, bodies, joints\n"
    "│   └── solver/              ← Generalized-α, Newton-Raphson\n"
    "│\n"
    "├── chassis_sim/             ← High-level chassis simulation\n"
    "│   ├── geometry.py\n"
    "│   ├── dynamics.py\n"
    "│   ├── sweep.py\n"
    "│   ├── stability.py\n"
    "│   └── tests/               ← Automated test suite\n"
    "│\n"
    "├── api/                     ← FastAPI REST server\n"
    "│   └── routers/             ← Endpoints for each module\n"
    "│\n"
    "├── graphify-out/            ← Knowledge graph outputs\n"
    "│   ├── graph.json           ← Full graph data (1.6 MB)\n"
    "│   ├── GRAPH_REPORT.md      ← Human-readable analysis\n"
    "│   └── mindmap.html         ← Interactive brain visualisation\n"
    "│\n"
    "├── brain_map.py             ← Script to regenerate the brain map\n"
    "└── build_graph.py           ← Script to rebuild the knowledge graph"
)

spacer()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. QUICK REFERENCE
# ═══════════════════════════════════════════════════════════════════════════════
h1("8. Quick Reference — Useful Commands")

h2("Open the brain mind map")
code_block("brainmap")

h2("Rebuild the mind map after code changes")
code_block(
    "cd /home/dikshant/Desktop/Moter_bike\n"
    "python3.13 -m graphify update .\n"
    "python3.13 brain_map.py\n"
    "brainmap"
)

h2("Use the Python physics engine")
code_block(
    "cd /home/dikshant/Desktop/Moter_bike\n"
    "python3.13\n\n"
    ">>> from dynamics_engine import MotorcycleDynamicsModel\n"
    ">>> m = MotorcycleDynamicsModel(preset='sport')\n"
    ">>> m.set_input('swingarm_length', 620)\n"
    ">>> m.print_report()\n"
    ">>> m.plot_sweep('swingarm_length', 480, 700, steps=30)"
)

h2("Run the test suite")
code_block(
    "cd /home/dikshant/Desktop/Moter_bike\n"
    "python3.13 -m pytest chassis_sim/tests/ -v"
)

h2("Ask the knowledge graph a question")
code_block(
    'python3.13 -m graphify query "how does anti-squat work"\n'
    'python3.13 -m graphify explain "MotorcycleDynamicsModel"\n'
    'python3.13 -m graphify path "swingarm_length" "anti_squat_pct"'
)

spacer()

# ═══════════════════════════════════════════════════════════════════════════════
# 9. KEY TERMS GLOSSARY
# ═══════════════════════════════════════════════════════════════════════════════
h1("9. Simple Glossary")
terms = [
    ("Trail",
     "The horizontal distance between where the front wheel touches the ground "
     "and where the steering axis hits the ground. More trail = more stable but "
     "heavier steering."),
    ("Rake (Head Angle)",
     "How far the front forks are angled back from vertical. A sport bike is "
     "around 24°, a cruiser around 32°."),
    ("Anti-Squat",
     "How much the bike's own chain force resists the rear squatting down under "
     "acceleration. 100% = perfectly balanced. Sport bikes aim for ~85–110%."),
    ("Anti-Dive",
     "How much the front suspension resists diving under braking. Linked to "
     "brake caliper mounting geometry."),
    ("Instant Centre (IC)",
     "An imaginary point in space where the swingarm line and chain force line "
     "cross. Its position determines anti-squat percentage."),
    ("Centre of Gravity (CoG)",
     "The single point where all the bike's weight effectively acts. Lower and "
     "more central is generally better for handling."),
    ("DAG",
     "Directed Acyclic Graph. A network of nodes where each node depends on "
     "others. Used here so changing one parameter auto-updates all downstream values."),
    ("MBD",
     "Multi-Body Dynamics. Simulating multiple connected rigid parts (bodies) "
     "with Newton's laws applied to each one simultaneously."),
    ("Generalized-α Solver",
     "A professional numerical integration method for MBD. More accurate and "
     "stable than simple Euler or RK4 for stiff mechanical systems."),
    ("Weave / Wobble / Capsize",
     "The three natural motion modes of a motorcycle. Weave is a slow "
     "side-to-side sway, wobble is fast handlebar oscillation, capsize is "
     "slow falling over. Stability analysis shows which speeds are dangerous."),
    ("Pacejka Magic Formula",
     "A standard mathematical model for tyre forces. Produces realistic grip "
     "curves used in all professional vehicle simulation software."),
    ("Knowledge Graph",
     "A network map where nodes are code entities (functions, classes) and "
     "edges show relationships (calls, imports, uses). Lets you navigate "
     "and query the codebase like a database."),
]
for term, definition in terms:
    h3(f"  {term}")
    p = doc.add_paragraph(definition)
    p.paragraph_format.left_indent = Inches(0.3)
    if p.runs:
        p.runs[0].font.size = Pt(11)

spacer()

# ═══════════════════════════════════════════════════════════════════════════════
# 10. WHAT'S NEXT
# ═══════════════════════════════════════════════════════════════════════════════
h1("10. What Could Come Next")
next_steps = [
    ("Contact & Collision (MBD Phase 4)",
     "Add wheel-ground contact forces so the MBD solver can simulate full "
     "suspension travel realistically."),
    ("Pacejka Tyre Integration",
     "Connect the Magic Formula tyre model to the dynamics simulation so grip "
     "limits and slide angles are computed."),
    ("3D Stability Analysis",
     "Extend the linearised stability from the 2D Meijaard model to the full "
     "3D MBD system."),
    ("Export to CAD",
     "Allow geometry to be exported as a DXF or STEP file for use in SolidWorks "
     "or Fusion 360."),
    ("Live Telemetry Import",
     "Read GPS/IMU data from a real bike and compare against simulation predictions."),
]
for title, desc in next_steps:
    bullet(f"{title}: {desc}", f"{title}:")

spacer()
divider()

end_p = doc.add_paragraph()
end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
er = end_p.add_run(
    "Built with Python · TypeScript · React · D3.js · FastAPI · graphify  |  "
    f"{datetime.date.today().year}"
)
er.font.size = Pt(10)
er.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
er.font.italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
out = "Motorcycle_Chassis_Workbench_Project_Report.docx"
doc.save(out)
print(f"Saved: {out}")
